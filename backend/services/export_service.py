"""Export service for Excel and Word documents."""
import io
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.dml.color import RGBColor
from sqlalchemy.orm import Session
from typing import Optional
import os

DARK_BLUE = RGBColor(0, 51, 102)


def _set_margins_a4_moderate(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)


def _set_default_font_ariel(doc):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    rFonts.set(qn("w:eastAsia"), "Arial")


def _style_run(run, size_pt=None, bold=None, color=None, font_name="Arial"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_header_image(doc, image_path):
    """Add header image to document."""
    if not image_path or not os.path.exists(image_path):
        return
    section = doc.sections[0]
    header = section.header
    for p in header.paragraphs:
        try:
            p.clear()
        except Exception:
            for rr in p.runs:
                rr.text = ""
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Cm(17.0))


class ExportService:
    """Service for exporting data to Excel and Word."""
    
    def __init__(self, db: Session):
        self.db = db
    
    def export_to_excel(self, calculation_id: Optional[int] = None) -> io.BytesIO:
        """Export calculation results to Excel."""
        from models import InvoiceItem, WarehouseStock, Material
        from services.calculation_service import CalculationService
        
        buffer = io.BytesIO()
        
        # Get all invoice items
        items = self.db.query(InvoiceItem).all()
        
        # Prepare invoice items data
        items_data = []
        for item in items:
            items_data.append({
                "Broj računa": item.invoice.invoice_number if item.invoice else "",
                "Materijal": item.material_name,
                "Količina za fakturisanje": item.quantity_for_billing,
                "Jedinica za fakturisanje": item.unit_for_billing,
                "Količina za skidanje sa lagera": item.quantity_for_warehouse,
                "Jedinica za lager": item.unit_for_warehouse,
                "Količina sa koef": item.quantity_with_coef,
                "Napomena konverzije": item.conversion_note,
            })
        
        df_items = pd.DataFrame(items_data)
        
        # Get comparison data
        service = CalculationService(self.db)
        comparison = service.calculate_comparison(0)
        
        df_comparison = pd.DataFrame(comparison.get("results", []))
        
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            if not df_comparison.empty:
                df_comparison.to_excel(writer, index=False, sheet_name="uporedba")
            if not df_items.empty:
                df_items.to_excel(writer, index=False, sheet_name="fakture_obracun")
        
        buffer.seek(0)
        return buffer
    
    def export_word_for_invoice(self, invoice_number: str) -> io.BytesIO:
        """Export Word document for a specific invoice."""
        from models import Invoice, InvoiceItem
        
        invoice = self.db.query(Invoice).filter(
            Invoice.invoice_number == invoice_number
        ).first()
        
        if not invoice:
            raise ValueError(f"Invoice {invoice_number} not found")
        
        items = self.db.query(InvoiceItem).filter(
            InvoiceItem.invoice_id == invoice.id
        ).all()
        
        doc = Document()
        _set_margins_a4_moderate(doc)
        _set_default_font_ariel(doc)
        
        # Add header image
        # Try multiple possible paths
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        header_img_path = os.path.join(base_dir, "header.png")
        if not os.path.exists(header_img_path):
            # Try relative to backend directory
            header_img_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "..", "header.png")
        _add_header_image(doc, header_img_path)
        
        doc.add_paragraph("")
        
        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("NORMATIV POTROŠNJE MATERIJALA ZA IZVOĐENJE HIDROIZOLACIJE")
        _style_run(r, size_pt=16, bold=True, color=DARK_BLUE)
        
        doc.add_paragraph("")
        
        # Invoice number
        p1 = doc.add_paragraph()
        r1 = p1.add_run(f"PO RAČUNU BROJ: {invoice_number}")
        _style_run(r1, bold=True)
        
        # Company
        p2 = doc.add_paragraph()
        r2a = p2.add_run("KLIJENT:\t\t     ")
        _style_run(r2a, bold=True)
        r2b = p2.add_run(invoice.company or "")
        _style_run(r2b, bold=True)
        
        doc.add_paragraph("")
        doc.add_paragraph(
            "Prema normativima proizvođača materijala za hidroizolaciju po sistemu predviđena je okvirna sledeća potrošnja materijala:"
        )
        
        # Materials table
        unique_materials = {}
        for item in items:
            if item.material_name not in unique_materials:
                unique_materials[item.material_name] = {
                    "description": item.material_description or "",
                    "tech_spec": item.technical_spec or "",
                }
        
        if unique_materials:
            t1 = doc.add_table(rows=len(unique_materials) + 1, cols=3)
            t1.style = "Table Grid"
            t1.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            headers = ["Materijal", "Opis Materijala", "Normativna potrošnja (tehnički list)"]
            for i, h in enumerate(headers):
                t1.cell(0, i).text = h
            
            for idx, (mat_name, info) in enumerate(unique_materials.items(), start=1):
                t1.cell(idx, 0).text = mat_name
                t1.cell(idx, 1).text = info["description"]
                t1.cell(idx, 2).text = info["tech_spec"]
        
        doc.add_paragraph("")
        doc.add_paragraph("NAPOMENA: Sve pomenute potrošnje su minimalne i mogu biti različite u zavisnosti od površine.")
        doc.add_paragraph("")
        
        # Second title
        p3 = doc.add_paragraph()
        r3 = p3.add_run("Stvarne potrosnje hidroizolacionog materijala za predmetni racun")
        _style_run(r3, size_pt=16, bold=True, color=DARK_BLUE)
        
        # Items table
        if items:
            headers = [
                "ID materijala",
                "Materijal",
                "Površina na koju je naneta – Fakturisana količina:",
                "Jedinica",
                "Stvarna potrosnja",
                "Jedinica",
            ]
            table = doc.add_table(rows=len(items) + 1, cols=len(headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for i, h in enumerate(headers):
                table.cell(0, i).text = h
            
            for idx, item in enumerate(items, start=1):
                table.cell(idx, 0).text = str(item.material_id_from_file or "")
                table.cell(idx, 1).text = item.material_name
                table.cell(idx, 2).text = str(item.quantity_normative or item.quantity_for_billing)
                table.cell(idx, 3).text = item.unit_normative or item.unit_for_billing
                table.cell(idx, 4).text = str(item.quantity_with_coef or item.quantity_for_warehouse or "")
                table.cell(idx, 5).text = item.unit_for_warehouse or ""
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer