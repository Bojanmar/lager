

"""obracun.py – logika obračuna zaliha + export (Excel/Word).

NAPOMENA (Jan 2026): U Word izveštaju po računu, donja tabela se više ne "spaja" (merge)
po materijalu. Svaki materijal ostaje u svom redu, sa sopstvenom fakturisanom količinom
i jedinicom. Takođe je dodata kolona "ID materijala" (iz IZLAZ fajla).
"""
import re
import io
import numpy as np
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from docx.shared import Pt, Cm
from docx.dml.color import RGBColor


DARK_BLUE = RGBColor(0, 51, 102)  # tamno plava

def _set_margins_a4_moderate(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

def _set_default_font_ariel(doc):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # bitno za Word da stvarno primeni font svuda
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    rFonts.set(qn("w:eastAsia"), "Arial")

def _style_run(run, size_pt=None, bold=None, color=None, font_name="Arial"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color

def _force_tables_font_ariel(doc, font_name="Arial"):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _style_run(r, font_name=font_name)

def _add_header_image(doc, image_path):
    """
    Ubacuje sliku u HEADER (fiksno na svakoj strani).
    Očekuje 'header.png' u istom folderu kao obracun.py (ili prosleđen put).
    """
    if not image_path or not os.path.exists(image_path):
        return

    section = doc.sections[0]
    header = section.header

    # očisti header sadržaj
    for p in header.paragraphs:
        try:
            p.clear()
        except Exception:
            # ako clear ne postoji u tvojoj verziji python-docx, ručno brišemo run-ove
            for rr in p.runs:
                rr.text = ""

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # širina po A4, uz moderate margine
    run.add_picture(image_path, width=Cm(17.0))

def _first_nonempty_value(series):
    try:
        s = series.dropna().astype(str).str.strip()
        s = s[s != ""]
        return s.iloc[0] if len(s) else ""
    except Exception:
        return ""


# ======================================================
# 0) Safe Excel read (Streamlit UploadedFile seek fix)
# ======================================================

def _seek0(f):
    try:
        f.seek(0)
    except Exception:
        pass

def read_excel_safe(xlsx_file, **kwargs):
    _seek0(xlsx_file)
    return pd.read_excel(xlsx_file, **kwargs)

# ======================================================
# 1) Normalizacija teksta / jedinica
# ======================================================

def norm_text(s):
    if pd.isna(s):
        return ""
    s = str(s).replace('"', ' ').replace("'", ' ')
    s = s.replace("\n", " ")
    return re.sub(r"\s+", " ", s.strip()).lower()

def norm_unit(u):
    u = norm_text(u)
    repl = {
        "m^2": "m2", "m²": "m2",
        "m^1": "m",  "m¹": "m", "m1": "m",
        "kom (rolni)": "kom", "kom (rolna)": "kom", "kom (pak)": "kom",
        "kg.": "kg",
        "l": "lit", "litar": "lit", "litra": "lit", "litara": "lit",
        "džak": "dzak", "djak": "dzak",
        "rolni": "rolna",
        "kom rupa": "kom",
    }
    return repl.get(u, u)

def _round_qty_by_unit(qty_series, unit_series):
    qty = pd.to_numeric(qty_series, errors="coerce")
    unit_norm = unit_series.map(norm_unit)
    return np.where(unit_norm.eq("kom"), qty.round(0), qty.round(1))

def _format_qty_for_output(val, unit):
    if pd.isna(val):
        return ""
    try:
        v = float(val)
    except Exception:
        return str(val)
    if norm_unit(unit) == "kom":
        return str(int(round(v)))
    return f"{round(v, 1):.1f}"

# ======================================================
# 2) Markup za iste jedinice
# ======================================================

same_unit_markup = {
    "borner multiplex av 4": 0.20,
    "volteco, volgrip h.1.10 light": 0.15,
    "vintex mp fr 1.5mm": 0.10,
    "geotekstil 300gr-m2": 0.10,
    "poliesterska tkanina filc 100%, rolna 2x1m": 0.10,
}

# ======================================================
# 3) Helper pravila – svaki rule nosi _meta + _rule_desc
# ======================================================

def _attach_meta(fn, rule_type, from_unit, to_unit, factor=None, extra=None, enabled=True):
    fn._meta = {
        "rule_type": rule_type,
        "from_unit": from_unit,
        "to_unit": to_unit,
        "factor": factor,
        "extra": extra,
        "enabled": enabled
    }
    return fn

def rule_factor_per(area_unit, out_unit, factor):
    def _f(q, uf, ul):
        if norm_unit(uf) != area_unit or norm_unit(ul) != out_unit:
            return None, f"pravilo_ne_važi({uf}->{ul})"
        return q * factor, f"factor {factor} {out_unit}/{area_unit}"
    return _attach_meta(_f, "factor_per", area_unit, out_unit, factor=factor)

def rule_factor_per_len(len_unit, out_unit, factor):
    def _f(q, uf, ul):
        if norm_unit(uf) != len_unit or norm_unit(ul) != out_unit:
            return None, f"pravilo_ne_važi({uf}->{ul})"
        return q * factor, f"factor {factor} {out_unit}/{len_unit}"
    return _attach_meta(_f, "factor_per_len", len_unit, out_unit, factor=factor)

def rule_per_piece(out_unit, factor):
    def _f(q, uf, ul):
        if norm_unit(uf) != "kom" or norm_unit(ul) != out_unit:
            return None, f"pravilo_ne_važi({uf}->{ul})"
        return q * factor, f"{factor} {out_unit}/kom"
    return _attach_meta(_f, "per_piece", "kom", out_unit, factor=factor)

def rule_identity(from_unit, to_unit):
    def _f(q, uf, ul):
        if norm_unit(uf) != from_unit or norm_unit(ul) != to_unit:
            return None, f"pravilo_ne_važi({uf}->{ul})"
        return q, f"1:1 {from_unit}->{to_unit}"
    return _attach_meta(_f, "identity", from_unit, to_unit, factor=None)

def rule_m2_to_rolna(m2_per_rolna, extra=0.0):
    def _f(q, uf, ul):
        if norm_unit(uf) != "m2" or norm_unit(ul) != "rolna":
            return None, f"pravilo_ne_važi({uf}->{ul})"
        rolls = (q / m2_per_rolna) * (1.0 + extra)
        return rolls, f"m2→rolna; {m2_per_rolna} m2/rolna; +{int(extra * 100)}%"
    return _attach_meta(_f, "m2_to_rolna", "m2", "rolna", factor=m2_per_rolna, extra=extra)

def rule_m_to_rolna(m_per_rolna, extra=0.0):
    def _f(q, uf, ul):
        if norm_unit(uf) != "m" or norm_unit(ul) != "rolna":
            return None, f"pravilo_ne_važi({uf}->{ul})"
        rolls = (q / m_per_rolna) * (1.0 + extra)
        return rolls, f"m→rolna; {m_per_rolna} m/rolna; +{int(extra * 100)}%"
    return _attach_meta(_f, "m_to_rolna", "m", "rolna", factor=m_per_rolna, extra=extra)

def rule_m2_to_lit(liters_per_m2):
    def _f(q, uf, ul):
        if norm_unit(uf) != "m2" or norm_unit(ul) != "lit":
            return None, f"pravilo_ne_važi({uf}->{ul})"
        return q * liters_per_m2, f"{liters_per_m2} lit/m2"
    return _attach_meta(_f, "m2_to_lit", "m2", "lit", factor=liters_per_m2)

def rule_m_to_lit(liters_per_m):
    def _f(q, uf, ul):
        if norm_unit(uf) != "m" or norm_unit(ul) != "lit":
            return None, f"pravilo_ne_važi({uf}->{ul})"
        return q * liters_per_m, f"{liters_per_m} lit/m"
    return _attach_meta(_f, "m_to_lit", "m", "lit", factor=liters_per_m)

def rule_kg_to_dzak(kg_per_dzak):
    def _f(q, uf, ul):
        if norm_unit(uf) != "kg" or norm_unit(ul) != "dzak":
            return None, f"pravilo_ne_važi({uf}->{ul})"
        return q / kg_per_dzak, f"1/{kg_per_dzak} dzak/kg"
    return _attach_meta(_f, "kg_to_dzak", "kg", "dzak", factor=kg_per_dzak)

# ======================================================
# 4) Pravila: material_rules + alias
# ======================================================

material_rules = {}
def extend_rules(name, rules_to_add):
    key = norm_text(name)
    material_rules.setdefault(key, []).extend(rules_to_add)

ALIASES = { norm_text("Auqa Smart DUR 2k, 4+4kg"): norm_text("alchimica aqua smart dur 2k") }
def canon_mat(name: str) -> str:
    k = norm_text(name)
    return ALIASES.get(k, k)

# --- osnovna pravila ---
extend_rules("alchimica aqua smart dur 2k", [rule_factor_per("m2", "kg", 0.20)])
extend_rules("hyperdesmo pb 2k a+b, 20+20lit", [rule_factor_per("m2", "kg", 3)])
extend_rules("alchimica water foam 1k lv", [rule_per_piece("kg", 0.3)])
extend_rules("waterfoam catalyst 1 kg", [rule_per_piece("kg", 0.015)])
extend_rules("aquasmart - pb 1k 10kg, kg", [rule_factor_per("m2", "kg", 0.8), rule_factor_per_len("m", "kg", 0.7)])
extend_rules("borner gebortol vs", [rule_factor_per("m2", "kg", 0.4)])
extend_rules("cold cure polyurea 2k  a+b", [rule_factor_per("m2", "kg", 2.0)])
extend_rules("dual seal 15mil lg 8,92m2 rolna", [
    # m2 → rolna (postojeće)
    rule_m2_to_rolna(8.92, extra=0.20),

    # ✅ NOVO PRAVILO: m → rolna
    # 14 m = 1 rolna, +20% tolerancije
    rule_m_to_rolna(14, extra=0.20),
])


extend_rules("hydrobloc 575 integral - 1k pu resin elastic 6.5kg", [rule_per_piece("kg", 0.3)])
extend_rules("hydrocat 514 - highly active accelerator", [rule_per_piece("kg", 0.002)])
extend_rules("hydrobloc 510 - second-foam", [rule_per_piece("kg", 0.3)])

extend_rules("hyperdesmo -  ady-e 4lit", [rule_factor_per("m2", "kg", 0.15)])
extend_rules("hyperdesmo grey 1k 25 kg kanta", [rule_factor_per("m2", "kg", 2.5)])
extend_rules("hyperseal 2k f, 12kg", [rule_factor_per_len("m", "kg", 2.5)])
extend_rules("illbruck pu901 600ml, kom", [rule_per_piece("kg", 0.045)])
extend_rules("microsealer pu, 20kg", [rule_factor_per("m2", "kg", 0.4), rule_factor_per_len("m", "kg", 0.2)])

extend_rules("resin bau creck flex 2k a+b, 10+10.8kg", [rule_per_piece("kg", 0.3)])
extend_rules("resin bau hydrogum, 20kg", [rule_per_piece("kg", 0.3)])
extend_rules("resin bau water stopper 20kg comp a, 1,4kg comp b", [rule_per_piece("kg", 0.3)])
extend_rules("stopaq 2100 aquastop 0,53kg, kom", [rule_per_piece("kg", 0.003)])

extend_rules("vandex am 10, 20kg džak", [rule_factor_per("m3", "kg", 6.0)])
extend_rules("vandex bb 75, 25kg", [rule_factor_per("m2", "kg", 3.4)])
extend_rules("vandex cemelast liquid 9kg", [rule_factor_per("m2", "kg", 3.4)])
extend_rules("vandex injection mortar (vim) 25kg", [rule_per_piece("kg", 0.11)])
extend_rules("vandex plug, 15kg kanta", [])
extend_rules("vandex uni moratar 1z 25kg", [rule_factor_per_len("m", "kg", 5.0), rule_per_piece("kg", 0.5)])
extend_rules("yapseal 106, komp a, 20kg", [rule_factor_per("m2", "kg", 3.5)])
extend_rules("yapseal 106, komp b, 10kg", [rule_factor_per("m2", "kg", 1.75)])
extend_rules("yapseal 106, komp a + b, 20+10kg", [rule_factor_per("m2", "kg", 3.0)])

# ======================================================
# 5) Konstante + dodatna pravila
# ======================================================

PB2K_DENSITY_KG_PER_L = 1.15
PB2K_LIT_PER_M2 = round(2.4 / PB2K_DENSITY_KG_PER_L, 3)  # ≈ 2.087

ME508_3x25m_ROLL_LENGTH_M = 25
ME508_4x50m_ROLL_LENGTH_M = 50
CEMFLEX_PLATE_M_PER_KOM = 2
OMEGA_HOLDERS_PER_PAK = 100
OMEGA_HOLDERS_PER_M = 5
OMEGA_PAK_PER_M = OMEGA_HOLDERS_PER_M / OMEGA_HOLDERS_PER_PAK
VANDEX_PLUG_KG_PER_KOM = 0.1
VOLGRIP_WIDTH_M = 0.10
DUR2K_KG_PER_KOM = 0.20
DUR2K_KG_PER_M = 0.20
DUALSEAL_STD_M_PER_ROLNA = 14
extend_rules("Cleaner, 5kg", [rule_per_piece("kg", 0.025)])

extend_rules("ILLBRUCK PU901 600ml, kom", [
    rule_factor_per_len("m", "kom", 0.07),
    rule_factor_per("m2", "kom", 0.05),
])

extend_rules("Illbruck ME 508 privremena traka 3x25m", [
    rule_factor_per("m2", "m", 1.0),
    rule_factor_per_len("m", "kom", 1.0 / ME508_3x25m_ROLL_LENGTH_M),
    rule_factor_per("m2", "kom", 1.0 / ME508_3x25m_ROLL_LENGTH_M),
])
extend_rules("Illbruck ME 508 privremena traka 4x50m", [
    rule_factor_per("m2", "m", 1.0),
    rule_factor_per_len("m", "kom", 1.0 / ME508_4x50m_ROLL_LENGTH_M),
    rule_factor_per("m2", "kom", 1.0 / ME508_4x50m_ROLL_LENGTH_M),
])

extend_rules("Auqa Smart DUR 2k, 4+4kg", [
    rule_factor_per("m2", "kg", 0.20),
    rule_factor_per_len("m", "kg", DUR2K_KG_PER_M),
    rule_per_piece("kg", DUR2K_KG_PER_KOM),
])

extend_rules("HYPERDESMO PB 2K A+B, 20+20lit", [
    rule_m2_to_lit(PB2K_LIT_PER_M2),
    rule_m_to_lit(PB2K_LIT_PER_M2),
])

extend_rules("Alumanation 301 FT, kanta", [rule_identity("kom", "kanta")])

extend_rules("CemFLEX VB Coated Steel Plate  15cm X 2m", [rule_per_piece("m", CEMFLEX_PLATE_M_PER_KOM)])

extend_rules("CemFLEX VB Omega holder  100kom pak", [
    rule_per_piece("pak", 1 / OMEGA_HOLDERS_PER_PAK),
    rule_factor_per_len("m", "pak", OMEGA_PAK_PER_M),
])

extend_rules("DUAL SEAL 15mil, STD 8,92m2 rolna", [
    rule_m2_to_rolna(8.92, extra=0.00),
    rule_m_to_rolna(DUALSEAL_STD_M_PER_ROLNA, extra=0.00),
])

extend_rules("HYPERDESMO 2K-W  Comp A+B, 1.5+7.5kg", [rule_factor_per("m2", "kg", 1.0)])

extend_rules("HYPERDESMO GREY 1k 25 kg kanta", [
    rule_per_piece("kg", 25.0),
    rule_factor_per_len("m", "kg", 2.5),
])

extend_rules("RESIN BAU Easy Inject, 20kg", [rule_per_piece("kg", 0.09)])

extend_rules("CONNECT KSKSEAL privremena traka", [
    rule_factor_per_len("m", "kom", 0.07),
    rule_factor_per("m2", "kom", 0.05),
])

extend_rules("VANDEX Injection Mortar (VIM) 25kg", [rule_per_piece("kg", 0.11)])

extend_rules("VANDEX PLUG, 15kg kanta", [rule_per_piece("kg", VANDEX_PLUG_KG_PER_KOM)])

extend_rules("VANDEX SUPER, (25kg)", [rule_factor_per("m2", "kg", 2.0)])

extend_rules("VOLTECO, Volgrip H.1.10 light", [rule_factor_per_len("m", "m2", VOLGRIP_WIDTH_M)])

_dual_par_key = norm_text("DUAL SEAL PARAGRANULAR, 23kg džak")
material_rules[_dual_par_key] = [
    rule_per_piece("kg", 23.0),
    rule_kg_to_dzak(23.0),
    rule_identity("kom", "dzak"),
]

# ======================================================
# 6) Injektiranje – trigger + kalibracija
# ======================================================

def _n(s): return norm_text(s)

INJEKT_TRIGGER = _n("Injektiranje aktivnih prodora")
PACKER_MAT_KEY = _n("ALU Packer 10/100 mm, kom")

PACKER_MULT_MIN, PACKER_MULT_MAX = 1.0, 1.8
RESIN_MULT_MIN, RESIN_MULT_MAX = 0.2, 3.0

def _broj_pakera_po_setu(df):
    brojevi = {}
    for racun, grp in df.groupby("Broj računa", dropna=False):
        g = grp.copy()
        g["_mat_norm"] = g["Materijal"].map(_n)
        g["_jm_norm"] = g["Jedinica mere za fakturisanje"].map(norm_unit)
        alu_mask = g["_mat_norm"].eq(PACKER_MAT_KEY)
        if alu_mask.any():
            base = pd.to_numeric(g.loc[alu_mask, "Količina za fakturisanje"], errors="coerce").max()
        else:
            base = pd.to_numeric(g.loc[g["_jm_norm"].eq("kom"), "Količina za fakturisanje"], errors="coerce").max()
        brojevi[_n(racun)] = base if pd.notna(base) else None
    return brojevi

def _override_injekt(row, mat_norm, uf_norm, ul_norm, broj_pakera_map):
    tip = row.get("Pozicija za fakturisanje - tip hidroizolacije")
    if pd.isna(tip) or INJEKT_TRIGGER not in _n(tip):
        return None, None
    return None, None

def _override_sanacija(row, mat, uf, ul):
    tip_hidro = row.get("Pozicija za fakturisanje - tip hidroizolacije")
    if pd.notna(tip_hidro) and norm_text(tip_hidro) == norm_text("Sanacija kapilarne vlage"):
        _uni_alias = {norm_text("VANDEX UNI MORTAR 1Z 25kg"), norm_text("VANDEX UNI MORATAR 1Z 25kg")}
        if mat in _uni_alias and uf == "m2" and ul == "kg":
            return row["Količina za fakturisanje"] * 4.0, "override(sanacija): 4 kg/m2"
    return None, None

# ======================================================
# 7) Heuristike
# ======================================================

pair_defaults = {("m2", "kg"): 1.5, ("m", "kg"): 0.30, ("kom", "kg"): 0.30}

def make_heur_rule(uf, ul):
    uf = norm_unit(uf); ul = norm_unit(ul)
    if (uf, ul) not in pair_defaults:
        return None
    f = pair_defaults[(uf, ul)]
    if (uf, ul) == ("m2", "kg"): return rule_factor_per("m2", "kg", f)
    if (uf, ul) == ("m", "kg"):  return rule_factor_per_len("m", "kg", f)
    if (uf, ul) == ("kom", "kg"):return rule_per_piece("kg", f)
    return None

def calc_skidanje(row, broj_pakera_map, rules_dict):
    mat_raw = row["Materijal"]
    qty = row["Količina za fakturisanje"]
    u_fakt = row["Jedinica mere za fakturisanje"]
    u_lager = row["Jedinica mere za lager - skidanje količine"]
    if pd.isna(mat_raw) or pd.isna(qty) or pd.isna(u_fakt) or pd.isna(u_lager):
        return pd.Series([pd.NA, "nedostaju_podaci"])

    mat = canon_mat(mat_raw)
    uf = norm_unit(u_fakt)
    ul = norm_unit(u_lager)

    out, note = _override_injekt(row, mat, uf, ul, broj_pakera_map)
    if out is not None:
        return pd.Series([out, note])

    out, note = _override_sanacija(row, mat, uf, ul)
    if out is not None:
        return pd.Series([out, note])

    # ✅ ključ: iste jedinice -> prepis (osim markup liste)
    if uf == ul:
        base = qty
        add = same_unit_markup.get(mat, 0.0)
        if add:
            return pd.Series([base * (1 + add), f"same_unit +{int(add * 100)}%"])
        return pd.Series([base, "same_unit"])

    rules = rules_dict.get(mat, [])
    if not rules:
        return pd.Series([pd.NA, f"nema_pravila({u_fakt}->{u_lager})"])

    for r in rules:
        meta = getattr(r, "_meta", {})
        if meta and meta.get("enabled") is False:
            continue
        out, note = r(qty, uf, ul)
        if out is not None:
            return pd.Series([out, f"rule: {note}"])

    return pd.Series([pd.NA, f"pravilo_ne_pokriva({u_fakt}->{u_lager})"])

# ======================================================
# 8) Rules DF (editor)
# ======================================================

RULE_TYPES = [
    "factor_per","factor_per_len","per_piece","identity",
    "m2_to_rolna","m_to_rolna","m2_to_lit","m_to_lit","kg_to_dzak",
]

def rules_to_df(rules_dict):
    rows = []
    for mat, rules in rules_dict.items():
        for r in rules:
            meta = getattr(r, "_meta", None)
            if not meta:
                continue
            rows.append({
                "Materijal": mat,
                "rule_type": meta.get("rule_type",""),
                "from_unit": meta.get("from_unit",""),
                "to_unit": meta.get("to_unit",""),
                "factor": meta.get("factor", None),
                "extra": meta.get("extra", 0.0),
                "enabled": meta.get("enabled", True),
            })
    return pd.DataFrame(rows)

def make_rule_from_row(row):
    rt = row["rule_type"]
    fu = row.get("from_unit", "")
    tu = row.get("to_unit", "")
    factor = row.get("factor", None)
    extra = row.get("extra", 0.0)
    enabled = bool(row.get("enabled", True))

    if rt == "factor_per":
        fn = rule_factor_per(fu, tu, float(factor))
    elif rt == "factor_per_len":
        fn = rule_factor_per_len(fu, tu, float(factor))
    elif rt == "per_piece":
        fn = rule_per_piece(tu, float(factor))
    elif rt == "identity":
        fn = rule_identity(fu, tu)
    elif rt == "m2_to_rolna":
        fn = rule_m2_to_rolna(float(factor), extra=float(extra))
    elif rt == "m_to_rolna":
        fn = rule_m_to_rolna(float(factor), extra=float(extra))
    elif rt == "m2_to_lit":
        fn = rule_m2_to_lit(float(factor))
    elif rt == "m_to_lit":
        fn = rule_m_to_lit(float(factor))
    elif rt == "kg_to_dzak":
        fn = rule_kg_to_dzak(float(factor))
    else:
        return None

    fn._meta["enabled"] = enabled
    return fn

def apply_rules_df(base_rules_dict, edited_df):
    if edited_df is None or edited_df.empty:
        return {k: v[:] for k, v in base_rules_dict.items()}

    df = edited_df.copy()
    for c in ["Materijal", "rule_type", "from_unit", "to_unit"]:
        df[c] = df[c].fillna("").astype(str)
    df["enabled"] = df["enabled"].fillna(True).astype(bool)
    df["extra"] = pd.to_numeric(df["extra"], errors="coerce").fillna(0.0)
    df["factor"] = pd.to_numeric(df["factor"], errors="coerce")

    out = {}
    for mat, grp in df.groupby("Materijal", dropna=False):
        rules = []
        for _, r in grp.iterrows():
            fn = make_rule_from_row(r.to_dict())
            if fn is not None:
                rules.append(fn)
        out[mat] = rules

    for mat, rules in base_rules_dict.items():
        if mat not in out:
            out[mat] = rules[:]
    return out

# ======================================================
# 9) Učitavanje “wide” lager/magacin excela
# ======================================================

def read_wide_stock_excel(xlsx_file, label="stanje"):
    df = read_excel_safe(xlsx_file)
    units = df.iloc[0].to_dict()
    values = df.iloc[1].to_dict()

    out = pd.DataFrame({
        "Materijal": list(df.columns),
        "Jedinica": [norm_unit(units.get(c)) for c in df.columns],
        label: [values.get(c) for c in df.columns],
    })
    out[label] = pd.to_numeric(out[label], errors="coerce")
    out["Materijal_key"] = out["Materijal"].map(norm_text)
    return out

# ======================================================
# 10) Injekt kalibracija
# ======================================================

def _distribute_with_caps(base_vals, target_sum, min_mult, max_mult):
    base = base_vals.fillna(0).astype(float).clip(lower=0.0)
    if base.sum() <= 0 or pd.isna(target_sum):
        return pd.Series(index=base.index, data=np.nan)

    mult = pd.Series(index=base.index, data=min_mult, dtype=float)
    cur = float((base * mult).sum())
    target = float(target_sum)

    if target <= cur:
        return mult

    need = target - cur
    cap_add = base * (max_mult - min_mult)
    cap_total = float(cap_add.sum())
    if cap_total <= 0:
        return mult

    if need >= cap_total:
        return pd.Series(index=base.index, data=max_mult, dtype=float)

    add = cap_add * (need / cap_total)
    add = add.clip(lower=0.0, upper=cap_add)
    mult = (min_mult + (add / base.replace(0, np.nan))).fillna(min_mult)
    mult = mult.clip(lower=min_mult, upper=max_mult)

    for _ in range(6):
        cur = float((base * mult).sum())
        err = target - cur
        if abs(err) < 1e-6:
            break
        if err > 0:
            free = (mult < max_mult) & (base > 0)
            if not free.any(): break
            cap = (base[free] * (max_mult - mult[free])).sum()
            if cap <= 0: break
            delta = base[free] * (err / cap)
            mult.loc[free] = (mult.loc[free] + delta).clip(upper=max_mult)
        else:
            free = (mult > min_mult) & (base > 0)
            if not free.any(): break
            cap = (base[free] * (mult[free] - min_mult)).sum()
            if cap <= 0: break
            delta = base[free] * ((-err) / cap)
            mult.loc[free] = (mult.loc[free] - delta).clip(lower=min_mult)

    return mult

def apply_injekt_packers_and_resins(df_posle, lager_long, mag_long=None):
    df = df_posle.copy()

    if mag_long is None:
        return df, pd.DataFrame()

    stock = pd.merge(
        lager_long[["Materijal_key", "Stanje_na_lageru"]],
        mag_long[["Materijal_key", "Stanje_na_magacinu"]],
        on="Materijal_key",
        how="left"
    )
    stock["Target_potrosnja"] = stock["Stanje_na_lageru"] - stock["Stanje_na_magacinu"]

    tip = df["Pozicija za fakturisanje - tip hidroizolacije"]
    mask_injekt = tip.fillna("").map(_n).str.contains(INJEKT_TRIGGER, na=False)

    df["_racun_key"] = df["Broj računa"].map(_n)
    df["_mat_key"] = df["Materijal"].map(_n)
    df["_jm_lager"] = df["Jedinica mere za lager - skidanje količine"].map(norm_unit)

    # --- 1) Pakere ---
    packer_lines = mask_injekt & df["_mat_key"].eq(PACKER_MAT_KEY)
    billed_packers = (
        df.loc[packer_lines]
        .groupby("_racun_key")["Količina za fakturisanje"]
        .max()
        .astype(float)
    )

    packer_target = stock.loc[stock["Materijal_key"].eq(PACKER_MAT_KEY), "Target_potrosnja"]
    packer_target = float(packer_target.iloc[0]) if len(packer_target) else np.nan

    if billed_packers.sum() > 0 and pd.notna(packer_target):
        packer_mult = _distribute_with_caps(
            billed_packers,
            target_sum=packer_target,
            min_mult=PACKER_MULT_MIN,
            max_mult=PACKER_MULT_MAX
        )
        real_packers = (billed_packers * packer_mult).round(0)
    else:
        packer_mult = pd.Series(index=billed_packers.index, data=np.nan, dtype=float)
        real_packers = billed_packers.copy()

    for racun_key, rp in real_packers.items():
        mask_r = packer_lines & df["_racun_key"].eq(racun_key)
        if mask_r.any():
            df.loc[mask_r, "Količina za skidanje sa lagera"] = float(rp)
            df.loc[mask_r, "Napomena konverzije"] = (
                "injekt_paker_kalibracija: "
                f"fakt={float(billed_packers.loc[racun_key]):.0f} "
                f"x {float(packer_mult.loc[racun_key]):.3f} -> {float(rp):.0f}"
            )

    # --- 2) Smole (kg) ---
    resin_lines = mask_injekt & df["_jm_lager"].eq("kg") & (~df["_mat_key"].eq(PACKER_MAT_KEY))

    df["Količina za skidanje sa lagera"] = pd.to_numeric(df["Količina za skidanje sa lagera"], errors="coerce")
    base_by_mat_racun = (
        df.loc[resin_lines]
        .groupby(["_mat_key", "_racun_key"])["Količina za skidanje sa lagera"]
        .sum(min_count=1)
        .fillna(0.0)
    )

    resin_targets = stock.set_index("Materijal_key")["Target_potrosnja"].to_dict()

    real_packers_w = real_packers.to_dict()
    def _w(rk):
        v = real_packers_w.get(rk, np.nan)
        return float(v) if pd.notna(v) and v > 0 else 1.0

    resin_mult_map = {}

    for mat_key in base_by_mat_racun.index.get_level_values(0).unique():
        target = resin_targets.get(mat_key, np.nan)
        base_sr = base_by_mat_racun.loc[mat_key].copy()

        if pd.isna(target) or target <= 0 or base_sr.sum() <= 0:
            for rk in base_sr.index:
                resin_mult_map[(mat_key, rk)] = 1.0
            continue

        weights = pd.Series({rk: _w(rk) for rk in base_sr.index}).astype(float)
        base_for_caps = base_sr * (weights / weights.mean())

        mult = _distribute_with_caps(
            base_for_caps,
            target_sum=target,
            min_mult=RESIN_MULT_MIN,
            max_mult=RESIN_MULT_MAX
        )
        mult = mult.clip(lower=RESIN_MULT_MIN, upper=RESIN_MULT_MAX)

        for rk in base_sr.index:
            resin_mult_map[(mat_key, rk)] = float(mult.loc[rk])

    for (mat_key, racun_key), m in resin_mult_map.items():
        mask = resin_lines & df["_mat_key"].eq(mat_key) & df["_racun_key"].eq(racun_key)
        if not mask.any():
            continue
        df.loc[mask, "Količina za skidanje sa lagera"] = df.loc[mask, "Količina za skidanje sa lagera"] * m
        df.loc[mask, "Napomena konverzije"] = df.loc[mask, "Napomena konverzije"].fillna("") + (
            f" | injekt_smola_kalibracija x{m:.3f}"
        )

    injekt_debug = pd.DataFrame({
        "Broj_racuna_key": billed_packers.index,
        "Fakturisano_pakera": billed_packers.values,
        "Koef_pakera": packer_mult.reindex(billed_packers.index).values,
        "Realno_pakera": real_packers.reindex(billed_packers.index).values,
    }).sort_values("Broj_racuna_key")

    return df, injekt_debug

# ======================================================
# 11) Glavna funkcija – procesiraj_obracun
# ======================================================

def procesiraj_obracun(lager_file, fakture_file, magacin_file=None, edited_rules_df=None, manual_map_df=None):
    """
    Vraća:
      - uporedba
      - df_posle
      - rules_used_df
      - kalibracija_ekstremi
      - injekt_debug
      - audit_map (IZLAZ↔LAGER)
      - sus_bad (JM iste a koef != 1)
    """
    df_fakture = read_excel_safe(fakture_file)

        # ======================================================
    # ✅ NOVO: obezbedi da "Normative" kolone postoje (za Word)
    # (ne utiče na obračun – samo se prenose kao info)
    # ======================================================
    col_norm_qty = "Količina za fakturisanje (ono što piše u tabeli za račune - Normative)"
    col_norm_jm_dash = "Jedinica mere za fakturisanje – u računu"  # en-dash
    col_norm_jm_hyph = "Jedinica mere za fakturisanje - u računu"  # običan minus

    # količina (normativi)
    if col_norm_qty not in df_fakture.columns:
        df_fakture[col_norm_qty] = pd.NA

    # jedinica (normativi) – prihvati oba naziva kolone
    has_dash = col_norm_jm_dash in df_fakture.columns
    has_hyph = col_norm_jm_hyph in df_fakture.columns
    if not has_dash and not has_hyph:
        # napravi obe da Word uvek ima šta da čita
        df_fakture[col_norm_jm_dash] = pd.NA
        df_fakture[col_norm_jm_hyph] = pd.NA
    elif has_dash and not has_hyph:
        df_fakture[col_norm_jm_hyph] = df_fakture[col_norm_jm_dash]
    elif has_hyph and not has_dash:
        df_fakture[col_norm_jm_dash] = df_fakture[col_norm_jm_hyph]


    # ✅ primeni ručno mapiranje (IZLAZ -> LAGER)
    if manual_map_df is not None and not manual_map_df.empty:
        mm = manual_map_df.copy()
        for c in ["Materijal_izlaz", "Materijal_lager"]:
            if c not in mm.columns:
                mm = pd.DataFrame(columns=["Materijal_izlaz", "Materijal_lager"])
                break
        if not mm.empty:
            mm = mm.dropna(subset=["Materijal_izlaz", "Materijal_lager"])
            mm["Materijal_izlaz"] = mm["Materijal_izlaz"].astype(str)
            mm["Materijal_lager"] = mm["Materijal_lager"].astype(str)
            mm = mm[mm["Materijal_lager"].str.strip() != ""]
            map_dict = dict(zip(mm["Materijal_izlaz"], mm["Materijal_lager"]))
            if map_dict:
                df_fakture["Materijal_original"] = df_fakture["Materijal"]
                df_fakture["Materijal"] = df_fakture["Materijal"].map(lambda x: map_dict.get(str(x), x))

    lager_long = read_wide_stock_excel(lager_file, label="Stanje_na_lageru")
    rules_dict = apply_rules_df(material_rules, edited_rules_df)

    # 1) mapiranje JM iz lagera na fakture (NORMALIZOVANO)
    df_uvoz = read_excel_safe(lager_file)
    mapa_jedinica_raw = df_uvoz.iloc[0].to_dict()
    mapa_jedinica_norm = {norm_text(mat): val for mat, val in mapa_jedinica_raw.items()}

    df_fakture["Jedinica mere za lager - skidanje količine"] = df_fakture["Materijal"].map(
        lambda x: mapa_jedinica_norm.get(norm_text(x), np.nan)
    )

    # 2) override JM spec
    mask_dual = df_fakture["Materijal"].map(norm_text).eq(_dual_par_key)
    df_fakture.loc[mask_dual, "Jedinica mere za lager - skidanje količine"] = "dzak"

    me508_3 = norm_text("Illbruck ME 508 privremena traka 3x25m")
    me508_4 = norm_text("Illbruck ME 508 privremena traka 4x50m")
    mask_me3 = df_fakture["Materijal"].map(norm_text).eq(me508_3)
    mask_me4 = df_fakture["Materijal"].map(norm_text).eq(me508_4)
    df_fakture.loc[mask_me3 | mask_me4, "Jedinica mere za lager - skidanje količine"] = "kom"

    broj_pakera_map = _broj_pakera_po_setu(df_fakture)

    # 3) prvi prolaz
    df_pre = df_fakture.copy()
    df_pre[["Količina za skidanje sa lagera", "Napomena konverzije"]] = df_pre.apply(
        lambda row: calc_skidanje(row, broj_pakera_map, rules_dict),
        axis=1
    )
    neuspeh_pre = df_pre[df_pre["Količina za skidanje sa lagera"].isna()].copy()

    # 4) heuristike
    rules_all = {k: v[:] for k, v in rules_dict.items()}
    for _, r in neuspeh_pre[["Materijal", "Jedinica mere za fakturisanje", "Jedinica mere za lager - skidanje količine"]].drop_duplicates().iterrows():
        rr = make_heur_rule(r["Jedinica mere za fakturisanje"], r["Jedinica mere za lager - skidanje količine"])
        if rr is not None:
            rules_all.setdefault(canon_mat(r["Materijal"]), []).append(rr)

    # 5) drugi prolaz
    df_posle = df_fakture.copy()
    df_posle[["Količina za skidanje sa lagera", "Napomena konverzije"]] = df_posle.apply(
        lambda row: calc_skidanje(row, broj_pakera_map, rules_all),
        axis=1
    )

    # 6) PB2K spec
    PB2K_NAME_NORM = norm_text("HYPERDESMO PB 2K A+B, 20+20lit")
    materijal_norm = df_posle["Materijal"].map(norm_text)
    jm_fakt_norm = df_posle["Jedinica mere za fakturisanje"].map(norm_unit)
    jm_lager_norm = df_posle["Jedinica mere za lager - skidanje količine"].map(norm_unit)

    mask_pb2k = materijal_norm.eq(PB2K_NAME_NORM)
    mask_pb2k_m2 = mask_pb2k & jm_fakt_norm.eq("m2") & jm_lager_norm.eq("lit")
    df_posle.loc[mask_pb2k_m2, "Količina za skidanje sa lagera"] = (
        pd.to_numeric(df_posle.loc[mask_pb2k_m2, "Količina za fakturisanje"], errors="coerce") * PB2K_LIT_PER_M2
    )
    mask_pb2k_m = mask_pb2k & jm_fakt_norm.eq("m") & jm_lager_norm.eq("lit")
    df_posle.loc[mask_pb2k_m, "Količina za skidanje sa lagera"] = (
        pd.to_numeric(df_posle.loc[mask_pb2k_m, "Količina za fakturisanje"], errors="coerce") * PB2K_LIT_PER_M2
    )

    # kom -> round
    mask_kom_final = jm_lager_norm.eq("kom")
    df_posle.loc[mask_kom_final, "Količina za skidanje sa lagera"] = pd.to_numeric(
        df_posle.loc[mask_kom_final, "Količina za skidanje sa lagera"], errors="coerce"
    ).round(0)

    # 7) magacin + injekt kalibracija
    injekt_debug = pd.DataFrame()
    mag_long = None
    if magacin_file is not None:
        mag_long = read_wide_stock_excel(magacin_file, label="Stanje_na_magacinu")
        df_posle, injekt_debug = apply_injekt_packers_and_resins(df_posle, lager_long, mag_long=mag_long)

    # ======================================================
    # AUDIT 1: Mapiranje IZLAZ ↔ LAGER
    # ======================================================
    lager_cols = pd.DataFrame({"Materijal_lager": list(df_uvoz.columns)})
    lager_cols["Materijal_key"] = lager_cols["Materijal_lager"].map(norm_text)
    lager_cols["JM_lager"] = lager_cols["Materijal_key"].map(lambda k: norm_unit(mapa_jedinica_norm.get(k, np.nan)))

    audit_map = df_posle[[
        "Materijal",
        "Jedinica mere za fakturisanje",
        "Jedinica mere za lager - skidanje količine",
        "Količina za fakturisanje",
        "Količina za skidanje sa lagera",
        "Napomena konverzije"
    ]].copy()

    if "Materijal_original" in df_posle.columns:
        audit_map["Materijal_original"] = df_posle["Materijal_original"]

    audit_map["Materijal_key"] = audit_map["Materijal"].map(norm_text)
    audit_map = audit_map.merge(lager_cols, on="Materijal_key", how="left")

    audit_map["JM_fakt"] = audit_map["Jedinica mere za fakturisanje"].map(norm_unit)
    audit_map["JM_lager_mapirana"] = audit_map["Jedinica mere za lager - skidanje količine"].map(norm_unit)
    audit_map["Match_u_lageru"] = audit_map["Materijal_lager"].notna()
    audit_map["JM_iste"] = (audit_map["JM_fakt"] == audit_map["JM_lager_mapirana"]) & (audit_map["JM_fakt"] != "")

    # ======================================================
    # AUDIT 2: JM iste a koef != 1
    # ======================================================
    qf = pd.to_numeric(df_posle["Količina za fakturisanje"], errors="coerce")
    qs = pd.to_numeric(df_posle["Količina za skidanje sa lagera"], errors="coerce")
    df_posle["_uf"] = df_posle["Jedinica mere za fakturisanje"].map(norm_unit)
    df_posle["_ul"] = df_posle["Jedinica mere za lager - skidanje količine"].map(norm_unit)

    df_posle["Koef_konverzije"] = np.where(qf > 0, qs / qf, np.nan)

    sus_bad = df_posle[
        (df_posle["_uf"] == df_posle["_ul"]) &
        (df_posle["_uf"] != "") &
        (~df_posle["Koef_konverzije"].fillna(1.0).round(6).eq(1.0))
    ].copy()

    # 8) ukupna potrošnja
    df_posle["Količina za skidanje sa lagera"] = pd.to_numeric(df_posle["Količina za skidanje sa lagera"], errors="coerce")
    potrosnja = (
        df_posle.groupby("Materijal", dropna=False)["Količina za skidanje sa lagera"]
        .sum(min_count=1)
        .reset_index()
        .rename(columns={"Količina za skidanje sa lagera": "Ukupna_potrošnja"})
    )
    potrosnja["Materijal_key"] = potrosnja["Materijal"].map(norm_text)

    lager_stanje = lager_long[["Materijal", "Materijal_key", "Stanje_na_lageru"]].copy()

    uporedba = pd.merge(potrosnja, lager_stanje, on="Materijal_key", how="outer", suffixes=("", "_lager"))
    uporedba["Materijal"] = uporedba["Materijal"].fillna(uporedba.get("Materijal_lager"))
    if "Materijal_lager" in uporedba.columns:
        uporedba = uporedba.drop(columns=["Materijal_lager"])

    uporedba["Razlika_pre"] = uporedba["Stanje_na_lageru"] - uporedba["Ukupna_potrošnja"]

    if mag_long is not None:
        mag_only = mag_long[["Materijal_key", "Stanje_na_magacinu"]]
        uporedba = pd.merge(uporedba, mag_only, on="Materijal_key", how="left")
    else:
        uporedba["Stanje_na_magacinu"] = np.nan

    uporedba["Realna_potrosnja_po_magacinu"] = uporedba["Stanje_na_lageru"] - uporedba["Stanje_na_magacinu"]

    def _koef_novi(row):
        pot = row.get("Ukupna_potrošnja")
        real = row.get("Realna_potrosnja_po_magacinu")
        if pd.isna(pot) or pot == 0 or pd.isna(real):
            return np.nan
        return real / pot

    uporedba["Koef_novi"] = uporedba.apply(_koef_novi, axis=1)
    uporedba["Finalna_potrošnja"] = uporedba["Ukupna_potrošnja"] * uporedba["Koef_novi"]

    # ======================================================
    # NOVO: Koef_novi koji NE dira "Prodaja racuna" (po materijalu)
    # ======================================================
    tip_cols = [
        "Pozicija za fakturisanje - tip hidroizolacije",
        "Tip računa",
        "Tip racuna",
    ]
    fixed_labels = {"prodaja racuna", "prodaja materijala"}
    mask_fixed = pd.Series(False, index=df_posle.index)
    for c in tip_cols:
        if c in df_posle.columns:
            mask_fixed = mask_fixed | df_posle[c].fillna("").map(norm_text).isin(fixed_labels)

    df_posle["_mat_key"] = df_posle["Materijal"].map(norm_text)  # stabilan ključ za join
    qs2 = pd.to_numeric(df_posle["Količina za skidanje sa lagera"], errors="coerce")

    fixed_sum = (
        df_posle.loc[mask_fixed]
        .assign(_qs2=qs2[mask_fixed])
        .groupby("_mat_key")["_qs2"]
        .sum(min_count=1)
        .rename("Fixed_sum")
    )
    total_sum = (
        df_posle.assign(_qs2=qs2)
        .groupby("_mat_key")["_qs2"]
        .sum(min_count=1)
        .rename("Total_sum")
    )
    fixed_sum = fixed_sum.reindex(total_sum.index).fillna(0.0)
    adj_sum = (total_sum - fixed_sum).rename("Adj_sum")

    uporedba = uporedba.merge(
        fixed_sum.reset_index().rename(columns={"_mat_key": "Materijal_key"}),
        on="Materijal_key",
        how="left"
    )
    uporedba = uporedba.merge(
        adj_sum.reset_index().rename(columns={"_mat_key": "Materijal_key"}),
        on="Materijal_key",
        how="left"
    )
    uporedba["Fixed_sum"] = uporedba["Fixed_sum"].fillna(0.0)

    def _koef_novi_adjusted(row):
        target = row.get("Realna_potrosnja_po_magacinu")
        adj = row.get("Adj_sum")
        fixed = row.get("Fixed_sum")
        if pd.isna(target) or pd.isna(adj) or adj <= 0:
            return np.nan
        target_adj = target - (fixed if pd.notna(fixed) else 0.0)
        if target_adj <= 0:
            return 0.0
        return target_adj / adj

    uporedba["Koef_novi"] = uporedba.apply(_koef_novi_adjusted, axis=1)
    uporedba["Finalna_potrošnja"] = uporedba["Ukupna_potrošnja"] * uporedba["Koef_novi"]
    CAL_MIN, CAL_MAX = 0.3, 3.0
    uporedba["Kalibracija_status"] = np.where(
        uporedba["Koef_novi"].between(CAL_MIN, CAL_MAX, inclusive="both"),
        "ok",
        "EKSTREMNO"
    )
    kalibracija_ekstremi = uporedba[uporedba["Kalibracija_status"].eq("EKSTREMNO")].copy()

    koef_map = uporedba.set_index("Materijal_key")["Koef_novi"].to_dict()
    df_posle["_koef_novi_mat"] = df_posle["_mat_key"].map(koef_map)

    k2 = pd.to_numeric(df_posle["_koef_novi_mat"], errors="coerce")
    df_posle["Kolicina za skidanje sa uracunatim Koef_novi za ovaj materijal"] = np.where(
        mask_fixed,
        qs2,
        np.where(k2.notna(), qs2 * k2, qs2)  # ako nema Koef_novi, prepiši originalno skidanje
    )

    # ======================================================
    # Zaokruživanje: 1 decimal svuda, osim "kom" -> ceo broj
    # ======================================================
    if "Jedinica mere za lager - skidanje količine" in df_posle.columns:
        if "Količina za skidanje sa lagera" in df_posle.columns:
            df_posle["Količina za skidanje sa lagera"] = _round_qty_by_unit(
                df_posle["Količina za skidanje sa lagera"],
                df_posle["Jedinica mere za lager - skidanje količine"]
            )
        if "Kolicina za skidanje sa uracunatim Koef_novi za ovaj materijal" in df_posle.columns:
            df_posle["Kolicina za skidanje sa uracunatim Koef_novi za ovaj materijal"] = _round_qty_by_unit(
                df_posle["Kolicina za skidanje sa uracunatim Koef_novi za ovaj materijal"],
                df_posle["Jedinica mere za lager - skidanje količine"]
            )

    if "Jedinica mere za fakturisanje" in df_posle.columns and "Količina za fakturisanje" in df_posle.columns:
        df_posle["Količina za fakturisanje"] = _round_qty_by_unit(
            df_posle["Količina za fakturisanje"],
            df_posle["Jedinica mere za fakturisanje"]
        )

    # Normative kolone (Word)
    col_norm_qty = "Količina za fakturisanje (ono što piše u tabeli za račune - Normative)"
    col_norm_jm_dash = "Jedinica mere za fakturisanje – u računu"
    col_norm_jm_hyph = "Jedinica mere za fakturisanje - u računu"
    if col_norm_qty in df_posle.columns:
        if col_norm_jm_dash in df_posle.columns:
            df_posle[col_norm_qty] = _round_qty_by_unit(df_posle[col_norm_qty], df_posle[col_norm_jm_dash])
        elif col_norm_jm_hyph in df_posle.columns:
            df_posle[col_norm_qty] = _round_qty_by_unit(df_posle[col_norm_qty], df_posle[col_norm_jm_hyph])

    # Uporedba: sve numeričke na 1 decimalu
    num_cols = uporedba.select_dtypes(include=[np.number]).columns
    if len(num_cols) > 0:
        uporedba[num_cols] = uporedba[num_cols].round(1)

    rules_used_df = rules_to_df(rules_dict)

    return uporedba, df_posle, rules_used_df, kalibracija_ekstremi, injekt_debug, audit_map, sus_bad

# ======================================================
# 12) Export
# ======================================================

def export_to_excel(uporedba, df_fakture_posle, injekt_debug=None, audit_map=None, sus_bad=None):
    buffer = io.BytesIO()

    # ✅ filtriraj kolone iz fakture_obracun za export (isto kao UI)
    drop_cols = [
        "Napomena konverzije",
        "_racun_key", "_mat_key", "_jm_lager", "_uf", "_ul",
        "Koef_konverzije",
        "_koef_novi_mat",
    ]
    fakture_export = df_fakture_posle.drop(columns=drop_cols, errors="ignore")

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        uporedba.to_excel(writer, index=False, sheet_name="uporedba")
        fakture_export.to_excel(writer, index=False, sheet_name="fakture_obracun")
        if injekt_debug is not None and not injekt_debug.empty:
            injekt_debug.to_excel(writer, index=False, sheet_name="injekt_debug")
        if audit_map is not None and not audit_map.empty:
            audit_map.to_excel(writer, index=False, sheet_name="audit_map")
        if sus_bad is not None and not sus_bad.empty:
            sus_bad.to_excel(writer, index=False, sheet_name="audit_jm_same_bad")

    buffer.seek(0)
    return buffer


from docx import Document
from io import BytesIO
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _set_cell_center(cell, vertical=True):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1.0
    if vertical:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_cell_left(cell, vertical=True):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1.0
    if vertical:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _wrap_cell(cell):
    # Word wrap flag
    tcPr = cell._tc.get_or_add_tcPr()
    wrap = OxmlElement("w:wordWrap")
    wrap.set(qn("w:val"), "1")
    tcPr.append(wrap)

    # paragraf spacing (da nema ogromnih razmaka)
    for p in cell.paragraphs:
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1.0


def _set_table_fixed_layout(table):
    # isključi autofit
    table.autofit = False

    # fiksni layout u Word-u
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)


def _set_col_widths(table, widths_cm):
    # widths_cm: lista širina po kolonama (u cm)
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _format_cell_value(x):
    """Lep format vrednosti za upis u Word ćelije."""
    import numpy as np
    import pandas as pd

    if x is None or pd.isna(x):
        return ""

    # ako je broj
    try:
        v = float(x)
        if np.isfinite(v):
            if abs(v - round(v)) < 1e-9:
                return str(int(round(v)))
            return str(round(v, 3)).rstrip("0").rstrip(".")
    except Exception:
        pass

    # fallback string
    return str(x).strip()

def _format_cell_value(x):
    """Siguran prikaz broja (za Word)."""
    if x is None or pd.isna(x):
        return ""
    try:
        v = float(x)
        if v.is_integer():
            return str(int(v))
        return str(round(v, 2))
    except Exception:
        return str(x)

def _format_id_int(x):
    """ID u Word-u mora biti ceo broj (bez .0)."""
    if x is None or pd.isna(x):
        return ""
    try:
        return str(int(float(x)))
    except Exception:
        s = str(x).strip()
        return s if s.isdigit() else ""


def generate_word_for_racun(df_fakture_posle, broj_racuna):
    df = df_fakture_posle.copy()
    df = df[df["Broj računa"] == broj_racuna].copy()

    # -----------------------------
    # KLIJENT (iz kolone Kompanija)
    # -----------------------------
    client_name = ""
    if "Kompanija" in df.columns:
        client_name = _first_nonempty_value(df["Kompanija"])

    # -----------------------------
    # kolone (kao ranije)
    # -----------------------------
    fakt_col = "Količina za fakturisanje (ono što piše u tabeli za račune - Normative)"
    fakt_jm_col = "Jedinica mere za fakturisanje - u računu"

    opis_col = "Opis Materijala"
    tech_col = "Normativna potrošnja (tehnički list)"

    # --- vrednosti iz "Normative" kolone za taj račun (spojeno u jednu ćeliju)
    vals = pd.to_numeric(df.get(fakt_col), errors="coerce") if fakt_col in df.columns else pd.Series([], dtype=float)
    vals = vals.dropna().tolist()

    jms = df.get(fakt_jm_col) if fakt_jm_col in df.columns else pd.Series([pd.NA] * len(df))
    jms = jms.dropna().astype(str).str.strip()
    jms = [x for x in jms.tolist() if x != ""]

    normative_text = "\n".join([str(int(v)) if float(v).is_integer() else str(v) for v in vals]) if len(vals) else ""

    jm_unique = []
    for u in jms:
        if u not in jm_unique:
            jm_unique.append(u)
    
    # ======================================================
    # Word prikaz za Tabelu 2:
    # - svaka stavka dobija (Normative količinu + Normative jedinicu) 1:1
    # - a zatim se (po potrebi) merge-uje po grupama (količina, jedinica),
    #   tako da ispadne kao u Excelu čak i kad redovi nisu bili grupisani.
    # ======================================================

    tip_col = "Pozicija za fakturisanje - tip hidroizolacije"
    fakt_col = "Količina za fakturisanje (ono što piše u tabeli za račune - Normative)"
    fakt_jm_col = "Jedinica mere za fakturisanje - u računu"

    df2 = df.copy()
    for c in [tip_col, fakt_col, fakt_jm_col]:
        if c not in df2.columns:
            df2[c] = pd.NA

    df2["__fakt_num__"] = pd.to_numeric(df2[fakt_col], errors="coerce")
    df2[tip_col] = df2[tip_col].astype("string")
    df2[fakt_jm_col] = df2[fakt_jm_col].astype("string")

    # Stabilan sort: grupiše po tipu radova pa po (količina, jedinica)
    df2 = df2.sort_values(
        by=[tip_col, "__fakt_num__", fakt_jm_col],
        kind="mergesort",
        na_position="last"
    ).drop(columns=["__fakt_num__"])

    # ======================================================
    # Word doc + HEADER slika + globalni stilovi
    # ======================================================
    doc = Document()

    # margine + Arial
    _set_margins_a4_moderate(doc)
    _set_default_font_ariel(doc)

    # header.png u istom folderu
    header_img_path = os.path.join(os.path.dirname(__file__), "header.png")
    _add_header_image(doc, header_img_path)

    # ------------------------------------------------------
    # FIKSNI TEKST + stilovi (zahtevi)
    # ------------------------------------------------------
    doc.add_paragraph("")  # razmak ispod headera

    # NASLOV (bold, centriran, Arial 16, tamno plav)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NORMATIV POTROŠNJE MATERIJALA ZA IZVOĐENJE HIDROIZOLACIJE")
    _style_run(r, size_pt=16, bold=True, color=DARK_BLUE)

    doc.add_paragraph("")

    # PO RAČUNU BROJ (bold)
    p1 = doc.add_paragraph()
    r1 = p1.add_run(f"PO RAČUNU BROJ: {broj_racuna}")
    _style_run(r1, bold=True)

    # KLIJENT (bold)
    p2 = doc.add_paragraph()
    r2a = p2.add_run("KLIJENT:\t\t     ")
    _style_run(r2a, bold=True)
    r2b = p2.add_run(f"{client_name}")
    _style_run(r2b, bold=True)

    doc.add_paragraph("")

    doc.add_paragraph(
        "Prema normativima proizvođača materijala za hidroizolaciju po sistemu predviđena je okvirna sledeća potrošnja materijala:"
    )

    # ======================================================
    # TABELA 1 (normativi / opis / tehnički list)
    # ======================================================
    cols_for_meta = ["Materijal", opis_col, tech_col]
    for c in cols_for_meta:
        if c not in df.columns:
            df[c] = pd.NA

    meta_df = df[cols_for_meta].copy()
    meta_df["__ord"] = range(len(meta_df))
    meta_df = meta_df.sort_values("__ord").drop(columns="__ord")
    meta_df = meta_df.drop_duplicates(subset=["Materijal"], keep="first")

    t1_headers = ["Materijal", "Opis Materijala", "Normativna potrošnja (tehnički list)"]
    t1 = doc.add_table(rows=len(meta_df) + 1, cols=3)
    t1.style = "Table Grid"
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER

    _set_table_fixed_layout(t1)
    _set_col_widths(t1, [4.0, 8.0, 5.5])

    for i, h in enumerate(t1_headers):
        cell = t1.cell(0, i)
        cell.text = h
        _set_cell_center(cell)
        _wrap_cell(cell)

    for r_idx, (_, row) in enumerate(meta_df.iterrows(), start=1):
        mat = "" if pd.isna(row.get("Materijal")) else str(row.get("Materijal"))
        opis = "" if pd.isna(row.get(opis_col)) else str(row.get(opis_col))
        tech = "" if pd.isna(row.get(tech_col)) else str(row.get(tech_col))

        c0 = t1.cell(r_idx, 0); c0.text = mat
        c1 = t1.cell(r_idx, 1); c1.text = opis
        c2 = t1.cell(r_idx, 2); c2.text = tech

        _set_cell_left(c0); _wrap_cell(c0)
        _set_cell_left(c1); _wrap_cell(c1)
        _set_cell_left(c2); _wrap_cell(c2)

    doc.add_paragraph("")
    doc.add_paragraph("NAPOMENA: Sve pomenute potrošnje su minimalne i mogu biti različite u zavisnosti od površine.")
    doc.add_paragraph("")

    # DRUGI NASLOV (bold, Arial 16, tamno plav)
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Stvarne potrosnje hidroizolacionog materijala za predmetni racun")
    _style_run(r3, size_pt=16, bold=True, color=DARK_BLUE)

    # ======================================================
    # TABELA 2 (stvarne potrošnje)
    # ======================================================
    # ✅ Donja tabela: SVAKI materijal ima svoj red (bez merge).
    # Uključujemo i "ID materijala" iz IZLAZ fajla.
    headers = [
        "ID materijala",
        "Materijal",
        "Površina na koju je naneta – Fakturisana količina:",
        "Jedinica",
        "Stvarna potrosnja",
        "Jedinica",
    ]

    # kolone potrebne za donju tabelu
    id_col = "ID materijala"
    if id_col not in df.columns:
        df[id_col] = pd.NA
    if fakt_col not in df.columns:
        df[fakt_col] = pd.NA
    if fakt_jm_col not in df.columns:
        # podrži i varijantu sa EN DASH (–)
        alt = "Jedinica mere za fakturisanje – u računu"
        if alt in df.columns:
            df[fakt_jm_col] = df[alt]
        else:
            df[fakt_jm_col] = pd.NA

    df2 = df[[
        id_col,
        "Materijal",
        fakt_col,
        fakt_jm_col,
        "Kolicina za skidanje sa uracunatim Koef_novi za ovaj materijal",
        "Jedinica mere za lager - skidanje količine",
    ]].copy()
    

    n_items = len(df2)
    table = doc.add_table(rows=n_items + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _set_table_fixed_layout(table)
    # ukupno 17.18 cm (CONTENT_WIDTH_CM)
    _set_col_widths(table, [2.30, 4.70, 3.40, 1.80, 3.30, 1.68])

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_center(cell)
        _wrap_cell(cell)

    for r_idx, (_, row) in enumerate(df2.iterrows(), start=1):
        # ID materijala
        c_id = table.rows[r_idx].cells[0]
        mid = row.get("ID materijala")
        c_id.text = _format_id_int(mid)

        _set_cell_center(c_id)
        _wrap_cell(c_id)

        # Materijal
        c_mat = table.rows[r_idx].cells[1]
        c_mat.text = "" if pd.isna(row.get("Materijal")) else str(row.get("Materijal"))
        _set_cell_left(c_mat)
        _wrap_cell(c_mat)

        
        # ✅ Normative količina (po redu)
        fq = row.get(fakt_col)
        fu = row.get(fakt_jm_col)
        c_fq = table.rows[r_idx].cells[2]
        c_fq.text = _format_qty_for_output(fq, fu)
        _set_cell_center(c_fq)
        _wrap_cell(c_fq)

        # ✅ Normative jedinica (po redu)
        fu = row.get(fakt_jm_col)
        c_fu = table.rows[r_idx].cells[3]
        c_fu.text = "" if pd.isna(fu) else str(fu)
        _set_cell_center(c_fu)
        _wrap_cell(c_fu)

        v = row.get("Kolicina za skidanje sa uracunatim Koef_novi za ovaj materijal")
        c_pot = table.rows[r_idx].cells[4]
        c_pot.text = _format_qty_for_output(v, jl)
        _set_cell_center(c_pot)
        _wrap_cell(c_pot)

        jl = row.get("Jedinica mere za lager - skidanje količine")
        c_jm = table.rows[r_idx].cells[5]
        c_jm.text = "" if pd.isna(jl) else str(jl)
        _set_cell_center(c_jm)
        _wrap_cell(c_jm)

    # (bez merge) – svaka stavka ostaje u svom redu
    # forsiraj Arial i u tabelama
    _force_tables_font_ariel(doc, "Arial")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
import zipfile

def generate_word_zip_all_racuni(df_fakture_posle):
    buffer = BytesIO()

    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for racun in sorted(df_fakture_posle["Broj računa"].dropna().unique()):
            doc_buf = generate_word_for_racun(df_fakture_posle, racun)
            zf.writestr(f"racun_{racun}.docx", doc_buf.getvalue())

    buffer.seek(0)
    return buffer

def generate_excel_for_racun(df_fakture_posle, broj_racuna):
    df = df_fakture_posle[df_fakture_posle["Broj računa"] == broj_racuna]

    cols = [
        "Materijal",
        "Količina za fakturisanje",
        "Jedinica mere za fakturisanje",
        "Kolicina za skidanje sa uracunatim Koef_novi za ovaj materijal",
        "Jedinica mere za lager - skidanje količine",
    ]
    df = df[cols]

    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=f"Racun_{broj_racuna}")

    buffer.seek(0)
    return buffer
def generate_excel_zip_all_racuni(df_fakture_posle):
    buffer = BytesIO()

    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for racun in sorted(df_fakture_posle["Broj računa"].dropna().unique()):
            xls_buf = generate_excel_for_racun(df_fakture_posle, racun)
            zf.writestr(f"racun_{racun}.xlsx", xls_buf.getvalue())

    buffer.seek(0)
    return buffer

def procesiraj_obracun_iz_db(df_items, edited_rules_df=None, manual_map_df=None):
    """
    df_items mora imati kolone:
    - Broj računa
    - Kompanija
    - ID materijala
    - Materijal
    - Količina za fakturisanje
    - Jedinica mere za fakturisanje
    - Jedinica mere za lager - skidanje količine   (ovo dolazi iz materials.uom)
    - Opis Materijala
    - Normativna potrošnja (tehnički list)
    - Količina za fakturisanje (ono što piše u tabeli za račune - Normative)  (može = količina)
    - Jedinica mere za fakturisanje - u računu (može = JM fakturisanja)
    """

    df_fakture = df_items.copy()

    # ručno mapiranje (ako ga koristiš)
    if manual_map_df is not None and not manual_map_df.empty:
        mm = manual_map_df.dropna(subset=["Materijal_izlaz", "Materijal_lager"]).copy()
        mm["Materijal_izlaz"] = mm["Materijal_izlaz"].astype(str)
        mm["Materijal_lager"] = mm["Materijal_lager"].astype(str)
        mm = mm[mm["Materijal_lager"].str.strip() != ""]
        map_dict = dict(zip(mm["Materijal_izlaz"], mm["Materijal_lager"]))
        if map_dict:
            df_fakture["Materijal_original"] = df_fakture["Materijal"]
            df_fakture["Materijal"] = df_fakture["Materijal"].map(lambda x: map_dict.get(str(x), x))

    # pravila
    rules_dict = apply_rules_df(material_rules, edited_rules_df)

    # broj pakera mapa (za injekt logiku)
    broj_pakera_map = _broj_pakera_po_setu(df_fakture)

    # obračun skidanja
    df_posle = df_fakture.copy()
    df_posle[["Količina za skidanje sa lagera", "Napomena konverzije"]] = df_posle.apply(
        lambda row: calc_skidanje(row, broj_pakera_map, rules_dict),
        axis=1
    )

    # za sada preskačemo “Koef_novi” godišnju kalibraciju (to je posebna funkcija kasnije)
    df_posle["Kolicina za skidanje sa uracunatim Koef_novi za ovaj materijal"] = pd.to_numeric(
        df_posle["Količina za skidanje sa lagera"], errors="coerce"
    )

    rules_used_df = rules_to_df(rules_dict)

    return df_posle, rules_used_df

