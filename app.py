import os
from io import BytesIO
from datetime import date, datetime

import pandas as pd
import streamlit as st
from sqlalchemy import create_engine, text

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import RGBColor

# PDF (preview)
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm

# Font for Serbian diacritics in ReportLab
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# PDF->Images (preview)
try:
    import fitz  # pymupdf
    HAVE_PYMUPDF = True
except Exception:
    HAVE_PYMUPDF = False

from obracun import (
    procesiraj_obracun_iz_db,
    rules_to_df,
    RULE_TYPES,
    material_rules,
    generate_word_for_racun,  # tvoj postojeći Word za obračun
)

# =========================
# CONFIG
# =========================
st.set_page_config(page_title="Lager Lux", layout="wide")
st.title("📦 Lager Lux – Računi + Obračun (DB)")

DB_URL = os.getenv("DB_URL", "postgresql+psycopg2://postgres:postgres@localhost:5432/lager_lux")
engine = create_engine(DB_URL, pool_pre_ping=True)

# Putanja do slike zaglavlja koju koristiš i u final Word-u za obračun.
# Stavi fajl npr. u: C:\Users\hp\Desktop\lager kalkulator Branko\assets\header.png
HEADER_IMG_PATH = os.getenv("HEADER_IMG_PATH", "assets/header.png")

SERBIAN_TTF_CANDIDATES = [
    os.getenv("SERBIAN_TTF_PATH", ""),  # ako ručno setuješ env
    "C:/Windows/Fonts/DejaVuSans.ttf",
    "C:/Windows/Fonts/Arial.ttf",
    "assets/DejaVuSans.ttf",
]

def _register_serbian_font():
    # ReportLab: registruj font koji podržava č/ć/đ/š/ž
    for p in SERBIAN_TTF_CANDIDATES:
        p = (p or "").strip()
        if p and os.path.exists(p):
            try:
                pdfmetrics.registerFont(TTFont("SERBIAN", p))
                return "SERBIAN"
            except Exception:
                continue
    # fallback: Helvetica (ne garantuje sva slova)
    return "Helvetica"

SERB_FONT_NAME = _register_serbian_font()

# =========================
# DB HELPERS
# =========================
def db_df(sql: str, params=None) -> pd.DataFrame:
    with engine.connect() as c:
        return pd.read_sql(text(sql), c, params=params or {})

def db_exec(sql: str, params=None) -> None:
    with engine.begin() as c:
        c.execute(text(sql), params or {})

# =========================
# DB MIGRATIONS (lightweight)
# =========================
def _ensure_invoice_columns():
    try:
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS payment_purpose TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS invoice_note TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS fx_rate DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS nbs_rate_services DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS nbs_rate_services_date TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS nbs_rate_advance DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS nbs_rate_advance_date TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS advance_invoice_id INTEGER")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS advance_applied_amount DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS total_due DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS situation_title TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS situation_period TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS object_desc TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS contract_value_eur DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS works_place TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS base_contract_no TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS base_contract_date TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS add_contract_no TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS add_contract_date TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS cumulative_current_eur DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS cumulative_prev_eur DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS monthly_current_eur DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS monthly_current_rsd DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS advance_deduction_eur DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS advance_deduction_rsd DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS total_due_eur DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS total_due_rsd DOUBLE PRECISION")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS vat_note_override TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS payment_currency_text TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS supervisor_name TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS receiver_name TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS receiver_title TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS supplier_name TEXT")
        db_exec("ALTER TABLE invoices ADD COLUMN IF NOT EXISTS supplier_title TEXT")
        db_exec("ALTER TABLE invoice_items ADD COLUMN IF NOT EXISTS item_note TEXT")
    except Exception:
        # ako DB ne dozvoljava ALTER (npr. prava), nastavi bez prekida
        pass

_ensure_invoice_columns()

def _get_or_create_advance_material_id() -> int:
    row = db_df(
        "SELECT material_id FROM materials WHERE LOWER(name)=LOWER(:n) LIMIT 1",
        {"n": "AVANS"}
    )
    if not row.empty:
        return int(row.iloc[0]["material_id"])
    with engine.begin() as c:
        mid = c.execute(
            text("INSERT INTO materials(name, uom) VALUES (:n, :u) RETURNING material_id"),
            {"n": "AVANS", "u": "kom"}
        ).scalar()
    return int(mid)

# =========================
# NUM HELPERS
# =========================
def _to_num(x, default=0.0) -> float:
    try:
        if x is None:
            return float(default)
        if pd.isna(x):
            return float(default)
        return float(x)
    except Exception:
        return float(default)

def _fmt_rs_money(x: float) -> str:
    # 12345.6 -> "12.345,60"
    s = f"{float(x):,.2f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")

def _fmt_qty(x: float) -> str:
    s = f"{float(x):,.2f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")
INVOICE_TYPES = ["AVANSI", "PRODAJA MATERIJALA", "USLUGE", "PRIVREMENA SITUACIJA"]

def _is_advance(t: str) -> bool:
    return str(t).strip().upper() in {"ADVANCE", "AVANSI"}

def _is_sale(t: str) -> bool:
    return str(t).strip().upper() in {"FINAL", "PRODAJA MATERIJALA"}

def _is_service(t: str) -> bool:
    return str(t).strip().upper() == "USLUGE"

def _is_temp_situation(t: str) -> bool:
    return str(t).strip().upper() == "PRIVREMENA SITUACIJA"

def build_service_lines(df_items: pd.DataFrame) -> pd.DataFrame:
    out = df_items.copy()
    out["name"] = out.get("name", "")
    out["item_note"] = out.get("item_note", "")
    out["uom"] = out.get("uom", "")
    out["unit_price"] = out.get("unit_price", 0).apply(lambda v: _to_num(v, 0.0))
    out["total_price"] = out.get("total_price", 0).apply(lambda v: _to_num(v, 0.0))
    # ako nije unet total_price, koristi unit_price
    out.loc[out["total_price"].isna() | (out["total_price"] == 0), "total_price"] = out["unit_price"]
    out["line_net"] = out["total_price"]
    out["line_vat"] = 0.0
    out["line_gross"] = out["line_net"]
    out["qty"] = 1.0
    return out


# =========================
# VAT / TOTALS HELPERS
# =========================
def calc_lines(df_items: pd.DataFrame, vat_rate: float) -> pd.DataFrame:
    out = df_items.copy()
    out["qty"] = out["qty"].apply(lambda v: _to_num(v, 0.0))
    out["unit_price"] = out.get("unit_price", 0).apply(lambda v: _to_num(v, 0.0))
    out["discount"] = out.get("discount", 0).apply(lambda v: _to_num(v, 0.0))

    out["line_net"] = (out["qty"] * out["unit_price"]) - out["discount"]
    out.loc[out["line_net"] < 0, "line_net"] = 0.0
    out["line_vat"] = out["line_net"] * float(vat_rate)
    out["line_gross"] = out["line_net"] + out["line_vat"]
    return out

def sum_totals(df_lines: pd.DataFrame) -> dict:
    return {
        "total_net": float(df_lines["line_net"].sum()) if not df_lines.empty else 0.0,
        "total_vat": float(df_lines["line_vat"].sum()) if not df_lines.empty else 0.0,
        "total_gross": float(df_lines["line_gross"].sum()) if not df_lines.empty else 0.0,
    }

# =========================
# WORD helpers
# =========================
def _set_cell_align(cell, h="center", v=True):
    for p in cell.paragraphs:
        if h == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif h == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if v:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def _set_run_blue(run):
    run.font.color.rgb = RGBColor(0, 0, 153)  # tamno plava (možeš menjati)

def _add_header_image(doc: Document):
    if HEADER_IMG_PATH and os.path.exists(HEADER_IMG_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(HEADER_IMG_PATH, width=Inches(6.5))

# =========================
# WORD: "RAČUN" (Invoice form 1:1 logika, bez template-a)
# =========================
def generate_invoice_word_from_db(invoice_id: int) -> BytesIO:
    inv = db_df("""
        SELECT
            i.invoice_id, i.invoice_no, i.invoice_type, i.status,
            i.issue_date, i.due_date, i.currency, COALESCE(i.vat_rate,0) AS vat_rate,
            COALESCE(i.total_net,0) AS total_net,
            COALESCE(i.total_vat,0) AS total_vat,
            COALESCE(i.total_gross,0) AS total_gross,
            i.delivery_note_no,              -- broj otpremnice (ako postoji)
            i.place_of_issue,                -- mesto izdavanja (ako postoji)
            i.place_of_supply,               -- mesto prometa (ako postoji)
            i.supply_date,                   -- datum prometa (ako postoji)
            i.payment_purpose,
            i.invoice_note,
            i.fx_rate,
            i.nbs_rate_services, i.nbs_rate_services_date,
            i.nbs_rate_advance, i.nbs_rate_advance_date,
            i.advance_applied_amount, i.invoice_type,
            i.situation_title, i.situation_period, i.object_desc,
            i.contract_value_eur, i.works_place,
            i.base_contract_no, i.base_contract_date,
            i.add_contract_no, i.add_contract_date,
            i.cumulative_current_eur, i.cumulative_prev_eur,
            i.monthly_current_eur, i.monthly_current_rsd,
            i.advance_deduction_eur, i.advance_deduction_rsd,
            i.total_due_eur, i.total_due_rsd,
            i.vat_note_override, i.payment_currency_text,
            i.supervisor_name, i.receiver_name, i.receiver_title,
            i.supplier_name, i.supplier_title,
            c.name AS client_name,
            c.address AS client_address,
            c.city AS client_city
        FROM invoices i
        JOIN clients c ON c.client_id = i.client_id
        WHERE i.invoice_id = :id
    """, {"id": invoice_id})
    if inv.empty:
        raise ValueError(f"Invoice not found: {invoice_id}")
    inv = inv.iloc[0].to_dict()

    items = db_df("""
        SELECT
            it.item_id,
            it.description,
            it.item_note,
            it.qty,
            it.uom,
            COALESCE(it.unit_price,0) AS unit_price,
            COALESCE(it.discount,0) AS discount,
            COALESCE(it.line_net,0) AS line_net
        FROM invoice_items it
        WHERE it.invoice_id = :id
        ORDER BY it.item_id
    """, {"id": invoice_id})

    # company settings (prodavac)
    cs = db_df("SELECT * FROM company_settings WHERE id=1")
    if cs.empty:
        db_exec("INSERT INTO company_settings(id, company_name) VALUES (1,'LASER - LUX d.o.o.')")
        cs = db_df("SELECT * FROM company_settings WHERE id=1")
    cs = cs.iloc[0].to_dict()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # header image (isto kao u final wordu)
    _add_header_image(doc)

    # Title line (plavo)
    p = doc.add_paragraph(cs.get("company_name") or "LASER - LUX d.o.o.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    _set_run_blue(run)

    doc.add_paragraph("")

    # 2 bloka: kupac (levo) + prodavac (desno)
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    left = t.cell(0, 0)
    right = t.cell(0, 1)

    # Kupac blok (1:1 logika)
    kupac_lines = []
    if inv.get("client_name"):
        kupac_lines.append(str(inv["client_name"]))
    if inv.get("client_address"):
        kupac_lines.append(str(inv["client_address"]))
    if inv.get("client_city"):
        kupac_lines.append(str(inv["client_city"]))
    left.text = "\n".join(kupac_lines) if kupac_lines else ""
    _set_cell_align(left, "left")

    # Prodavac blok (hardkodovan / company_settings)
    seller_lines = []
    if cs.get("address"): seller_lines.append(cs["address"])
    if cs.get("city"): seller_lines.append(cs["city"])
    if cs.get("country"): seller_lines.append(cs["country"])
    if cs.get("phone"): seller_lines.append(f"tel. {cs['phone']}")
    if cs.get("fax"): seller_lines.append(f"tel/fax: {cs['fax']}")
    if cs.get("email"): seller_lines.append(f"E-mail: {cs['email']}")
    if cs.get("bank_account"): seller_lines.append(f"Račun: {cs['bank_account']}")
    pib = cs.get("pib") or ""
    mb = cs.get("mb") or ""
    if pib or mb:
        seller_lines.append(f"PIB: {pib}    Matični broj: {mb}".strip())
    right.text = "\n".join(seller_lines)
    _set_cell_align(right, "left")

    doc.add_paragraph("")

    # PRIVREMENA SITUACIJA (poseban layout)
    if _is_temp_situation(inv.get("invoice_type")):
        title = doc.add_paragraph(f"{inv.get('situation_title') or 'PRIVREMENA SITUACIJA'} (RAČUN BR. {inv.get('invoice_no','')})")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.runs[0]
        r.bold = True
        r.font.size = Pt(14)

        doc.add_paragraph("")

        # Tabela 1
        t1 = doc.add_table(rows=8, cols=2)
        t1.style = "Table Grid"
        rows = [
            ("Radovi izvršeni do kraja meseca", inv.get("situation_period") or ""),
            ("Na objektu", inv.get("object_desc") or ""),
            ("Vrednost ugovorenih radova", f"{_fmt_rs_money(_to_num(inv.get('contract_value_eur'),0.0))} eur"),
            (inv.get("vat_note_override") or "PDV se ne obračunava na osnovu člana 10, stav 2, tačka 3, Zakona o PDVu.", "0"),
            ("Mesto gde se radovi izvode", inv.get("works_place") or ""),
            ("Vrednost ispostavljene situacije za naplatu bez PDVa", f"{_fmt_rs_money(_to_num(inv.get('monthly_current_eur'),0.0))} eur"),
            ("Broj osnovnog ugovora", f"{inv.get('base_contract_no') or ''} {('od ' + inv.get('base_contract_date')) if inv.get('base_contract_date') else ''}".strip()),
            ("Broj dopunskog ugovora", f"{inv.get('add_contract_no') or '/'} {('od ' + inv.get('add_contract_date')) if inv.get('add_contract_date') else ''}".strip()),
        ]
        for i, (l, v) in enumerate(rows):
            t1.cell(i, 0).text = l
            t1.cell(i, 1).text = v

        doc.add_paragraph("")

        # Tabela 2
        t2 = doc.add_table(rows=5, cols=3)
        t2.style = "Table Grid"
        t2.cell(0, 0).text = "1. KUMULATIV IZVRŠENIH RADOVA PO TEKUĆOJ SITUACIJI"
        t2.cell(0, 1).text = f"{_fmt_rs_money(_to_num(inv.get('cumulative_current_eur'),0.0))} eur"
        t2.cell(0, 2).text = ""

        t2.cell(1, 0).text = "2. KUMULATIV IZVRŠENIH RADOVA PO PRETHODNOJ SITUACIJI"
        t2.cell(1, 1).text = f"{_fmt_rs_money(_to_num(inv.get('cumulative_prev_eur'),0.0))} eur"
        t2.cell(1, 2).text = ""

        t2.cell(2, 0).text = "3. MESEČNO IZVRŠENI RADOVI PO TEKUĆOJ SITUACIJI (1-2)"
        t2.cell(2, 1).text = f"{_fmt_rs_money(_to_num(inv.get('monthly_current_eur'),0.0))} eur"
        t2.cell(2, 2).text = f"{_fmt_rs_money(_to_num(inv.get('monthly_current_rsd'),0.0))} din"

        t2.cell(3, 0).text = "3. UMANJENJE PO OSNOVU AVANSA"
        t2.cell(3, 1).text = f"{_fmt_rs_money(_to_num(inv.get('advance_deduction_eur'),0.0))} eur"
        t2.cell(3, 2).text = f"{_fmt_rs_money(_to_num(inv.get('advance_deduction_rsd'),0.0))} rsd"

        t2.cell(4, 0).text = "UKUPNO ZA NAPLATU PO OVOJ SITUACIJI:"
        t2.cell(4, 1).text = f"{_fmt_rs_money(_to_num(inv.get('total_due_eur'),0.0))} eur"
        t2.cell(4, 2).text = f"{_fmt_rs_money(_to_num(inv.get('total_due_rsd'),0.0))} rsd"

        doc.add_paragraph("")
        doc.add_paragraph(inv.get("vat_note_override") or "PDV se ne obračunava na osnovu člana 10, stav 2, tačka 3, Zakona o PDVu.")
        if inv.get("nbs_rate_services"):
            doc.add_paragraph(f"Za obračun radova primenjen srednji kurs NBS na dan {inv.get('nbs_rate_services_date') or ''}. {inv.get('nbs_rate_services')}")
        if inv.get("nbs_rate_advance"):
            doc.add_paragraph(f"Za obračun avansa primenjen srednji kurs NBS na dan {inv.get('nbs_rate_advance_date') or ''}. {inv.get('nbs_rate_advance')}")
        if inv.get("issue_date"):
            doc.add_paragraph(f"Datum izdavanja situacije: {inv.get('issue_date')}")
        if inv.get("supply_date"):
            doc.add_paragraph(f"Datum prometa dobara i usluga: {inv.get('supply_date')}")
        if inv.get("payment_currency_text"):
            doc.add_paragraph(f"Valuta plaćanja: {inv.get('payment_currency_text')}")

        doc.add_paragraph("")
        doc.add_paragraph("Nadzorni organ")
        sig = doc.add_table(rows=2, cols=2)
        sig.cell(0, 0).text = "Primalac dobara i usluga"
        sig.cell(0, 1).text = "Ispručilac dobara i usluga"
        sig.cell(1, 0).text = f"{inv.get('receiver_name') or ''}, {inv.get('receiver_title') or ''}".strip(", ")
        sig.cell(1, 1).text = f"{inv.get('supplier_name') or ''}, {inv.get('supplier_title') or ''}".strip(", ")

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio

    # meta info
    issue_date = inv.get("issue_date")
    due_date = inv.get("due_date")
    supply_date = inv.get("supply_date")
    place_issue = inv.get("place_of_issue") or "Beograd"
    place_supply = inv.get("place_of_supply") or "Beograd"
    delivery_note_no = inv.get("delivery_note_no")

    def dstr(d):
        if isinstance(d, (datetime, date)):
            return d.strftime("%d.%m.%Y.")
        return ""

    meta_lines = [
        f"Datum izdavanja računa: {dstr(issue_date)}",
        f"Mesto izdavanja računa: {place_issue}",
    ]
    if delivery_note_no:
        meta_lines.append(f"Broj otpremnice: {delivery_note_no}")
    if supply_date:
        meta_lines.append(f"Mesto i datum prometa dobara: {place_supply}, {dstr(supply_date)}")
    meta_lines.append(f"Rok za plaćanje: {dstr(due_date) if due_date else '/'}")

    for line in meta_lines:
        doc.add_paragraph(line)

    doc.add_paragraph("")

    # naslov RAČUN BR. (plavo)
    title = doc.add_paragraph(f"RAČUN  BR. {inv.get('invoice_no','')}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.runs[0]
    r.bold = True
    r.font.size = Pt(14)
    _set_run_blue(r)

    doc.add_paragraph("Za nabavku materijala i to:")

    # tabela stavki (kao na slici)
    tbl = doc.add_table(rows=1, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    hdr[0].text = "R.br."
    hdr[1].text = "Ime proizvoda"
    hdr[2].text = "JM"
    hdr[3].text = "Količina"
    hdr[4].text = "Jed.cena"
    hdr[5].text = "Ukupno"

    for c in hdr:
        _set_cell_align(c, "center")
        for p in c.paragraphs:
            for rr in p.runs:
                rr.bold = True

    currency = inv.get("currency") or "RSD"

    for i, rrow in enumerate(items.to_dict("records"), start=1):
        row = tbl.add_row().cells
        row[0].text = str(i)
        row[1].text = str(rrow.get("description") or "")
        row[2].text = str(rrow.get("uom") or "")
        row[3].text = _fmt_qty(_to_num(rrow.get("qty"), 0.0))
        row[4].text = _fmt_rs_money(_to_num(rrow.get("unit_price"), 0.0))
        row[5].text = f"{_fmt_rs_money(_to_num(rrow.get('line_net'), 0.0))} {currency}"

        _set_cell_align(row[0], "center")
        _set_cell_align(row[1], "left")
        _set_cell_align(row[2], "center")
        _set_cell_align(row[3], "right")
        _set_cell_align(row[4], "right")
        _set_cell_align(row[5], "right")

    doc.add_paragraph("")

    total_net = _to_num(inv.get("total_net"), 0.0)
    total_vat = _to_num(inv.get("total_vat"), 0.0)
    total_gross = _to_num(inv.get("total_gross"), 0.0)
    vat_rate = _to_num(inv.get("vat_rate"), 0.0)

    p1 = doc.add_paragraph(f"UKUPNO : {_fmt_rs_money(total_net)} {currency}")
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.runs[0].bold = True

    p2 = doc.add_paragraph(f"OSNOVICA ZA OBRAČUN PDV: {_fmt_rs_money(total_net)} {currency}")
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p3 = doc.add_paragraph(f"PDV {int(vat_rate*100)}%: {_fmt_rs_money(total_vat)} {currency}")
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p4 = doc.add_paragraph(f"SVE UKUPNO ZA UPLATU: {_fmt_rs_money(total_gross)} {currency}")
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4.runs[0].bold = True

    doc.add_paragraph("")
    doc.add_paragraph(f"Paritet isporuke: {cs.get('delivery_parity') or '/'}")
    doc.add_paragraph(f"Napomena o poreskom oslobađanju: {cs.get('tax_note') or '/'}")
    doc.add_paragraph(cs.get("payment_note") or f"Uplatu izvršiti na tekući račun br.: {cs.get('bank_account') or ''}")
    if inv.get("payment_purpose"):
        doc.add_paragraph(f"Svrha uplate: {inv.get('payment_purpose')}")
    if inv.get("invoice_note"):
        doc.add_paragraph(f"Napomena: {inv.get('invoice_note')}")

    doc.add_paragraph("")
    sig = doc.add_paragraph(f"■ {cs.get('company_name') or 'LASER LUX'} ■ d.o.o.")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# =========================
# PDF preview (ReportLab) + PDF->Images
# =========================
def build_invoice_pdf_preview(
    company_block: dict,
    client_block: dict,
    meta: dict,
    lines_df: pd.DataFrame,
    totals: dict,
    currency: str,
    vat_percent: float
) -> BytesIO:
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    c.setTitle("Racun Preview")

    # font
    c.setFont(SERB_FONT_NAME, 12)

    y = h - 20 * mm

    # header image (optional)
    if HEADER_IMG_PATH and os.path.exists(HEADER_IMG_PATH):
        try:
            c.drawImage(HEADER_IMG_PATH, 20 * mm, y - 18 * mm, width=170 * mm, height=18 * mm, preserveAspectRatio=True, mask='auto')
            y -= 22 * mm
        except Exception:
            pass

    # company title (blue-ish simulation: reportlab uses RGB)
    c.setFillColorRGB(0, 0, 0.6)
    c.setFont(SERB_FONT_NAME, 18)
    c.drawCentredString(w / 2, y, company_block.get("company_name", "LASER - LUX d.o.o."))
    c.setFillColorRGB(0, 0, 0)
    c.setFont(SERB_FONT_NAME, 12)
    y -= 14 * mm

    # blocks
    left_x = 20 * mm
    right_x = 110 * mm

    # client block
    client_lines = [client_block.get("name",""), client_block.get("address",""), client_block.get("city","")]
    client_lines = [x for x in client_lines if x]
    yy = y
    for line in client_lines:
        c.drawString(left_x, yy, str(line))
        yy -= 5 * mm

    # company block
    comp_lines = []
    for k in ["address","city","country"]:
        if company_block.get(k): comp_lines.append(company_block[k])
    if company_block.get("phone"): comp_lines.append(f"tel. {company_block['phone']}")
    if company_block.get("fax"): comp_lines.append(f"tel/fax: {company_block['fax']}")
    if company_block.get("email"): comp_lines.append(f"E-mail: {company_block['email']}")
    if company_block.get("bank_account"): comp_lines.append(f"Račun: {company_block['bank_account']}")
    pib = company_block.get("pib") or ""
    mb = company_block.get("mb") or ""
    if pib or mb:
        comp_lines.append(f"PIB: {pib}    Matični broj: {mb}".strip())

    yy2 = y
    for line in comp_lines:
        c.drawString(right_x, yy2, str(line))
        yy2 -= 5 * mm

    # PRIVREMENA SITUACIJA (preview)
    if str(meta.get("invoice_type","")).upper() == "PRIVREMENA SITUACIJA":
        y -= 10 * mm
        c.setFont(SERB_FONT_NAME, 14)
        c.drawCentredString(w / 2, y, f"{meta.get('situation_title') or 'PRIVREMENA SITUACIJA'} (RAČUN BR. {meta.get('invoice_no','')})")
        y -= 10 * mm
        c.setFont(SERB_FONT_NAME, 11)
        lines = [
            ("Radovi izvršeni do kraja meseca", meta.get("situation_period","")),
            ("Na objektu", meta.get("object_desc","")),
            ("Vrednost ugovorenih radova", f"{_fmt_rs_money(_to_num(meta.get('contract_value_eur'),0.0))} eur"),
            ("PDV", "0"),
            ("Mesto gde se radovi izvode", meta.get("works_place","")),
            ("Vrednost ispostavljene situacije bez PDVa", f"{_fmt_rs_money(_to_num(meta.get('monthly_current_eur'),0.0))} eur"),
            ("Broj osnovnog ugovora", meta.get("base_contract_no","")),
            ("Broj dopunskog ugovora", meta.get("add_contract_no","")),
        ]
        for label, val in lines:
            c.drawString(left_x, y, f"{label}: {val}")
            y -= 6 * mm
            if y < 40*mm:
                c.showPage()
                c.setFont(SERB_FONT_NAME, 11)
                y = h - 20*mm

        y -= 4 * mm
        c.drawString(left_x, y, f"1) Kumulativ tekuća: {_fmt_rs_money(_to_num(meta.get('cumulative_current_eur'),0.0))} eur")
        y -= 6 * mm
        c.drawString(left_x, y, f"2) Kumulativ prethodna: {_fmt_rs_money(_to_num(meta.get('cumulative_prev_eur'),0.0))} eur")
        y -= 6 * mm
        c.drawString(left_x, y, f"3) Mesečno: {_fmt_rs_money(_to_num(meta.get('monthly_current_eur'),0.0))} eur / {_fmt_rs_money(_to_num(meta.get('monthly_current_rsd'),0.0))} din")
        y -= 6 * mm
        c.drawString(left_x, y, f"Umanjenje avansa: {_fmt_rs_money(_to_num(meta.get('advance_deduction_eur'),0.0))} eur / {_fmt_rs_money(_to_num(meta.get('advance_deduction_rsd'),0.0))} rsd")
        y -= 6 * mm
        c.drawString(left_x, y, f"Ukupno za naplatu: {_fmt_rs_money(_to_num(meta.get('total_due_eur'),0.0))} eur / {_fmt_rs_money(_to_num(meta.get('total_due_rsd'),0.0))} rsd")

        y -= 8 * mm
        c.drawString(left_x, y, meta.get("vat_note_override") or "")
        y -= 6 * mm
        if meta.get("nbs_rate_services"):
            c.drawString(left_x, y, f"NBS (radovi): {meta.get('nbs_rate_services_date','')} {meta.get('nbs_rate_services')}")
            y -= 6 * mm
        if meta.get("nbs_rate_advance"):
            c.drawString(left_x, y, f"NBS (avans): {meta.get('nbs_rate_advance_date','')} {meta.get('nbs_rate_advance')}")
            y -= 6 * mm
        c.drawString(left_x, y, f"Datum izdavanja: {meta.get('issue_date_str','')}")
        y -= 6 * mm
        c.drawString(left_x, y, f"Datum prometa: {meta.get('supply_date_str','')}")
        y -= 6 * mm
        c.drawString(left_x, y, f"Valuta plaćanja: {meta.get('payment_currency_text') or ''}")

        c.showPage()
        c.save()
        buf.seek(0)
        return buf

    y = min(yy, yy2) - 8 * mm

    # meta
    meta_lines = [
        f"Datum izdavanja računa: {meta.get('issue_date_str','')}",
        f"Mesto izdavanja računa: {meta.get('place_of_issue','Beograd')}",
    ]
    if meta.get("delivery_note_no"):
        meta_lines.append(f"Broj otpremnice: {meta['delivery_note_no']}")
    if meta.get("supply_date_str"):
        meta_lines.append(f"Mesto i datum prometa dobara: {meta.get('place_of_supply','Beograd')}, {meta['supply_date_str']}")
    meta_lines.append(f"Rok za plaćanje: {meta.get('due_date_str','/')}")

    for line in meta_lines:
        c.drawString(left_x, y, line)
        y -= 6 * mm

    y -= 4 * mm

    # title
    c.setFillColorRGB(0, 0, 0.6)
    c.setFont(SERB_FONT_NAME, 16)
    c.drawCentredString(w / 2, y, f"RAČUN  BR. {meta.get('invoice_no','')}")
    c.setFillColorRGB(0, 0, 0)
    c.setFont(SERB_FONT_NAME, 12)
    y -= 10 * mm

    c.drawString(left_x, y, "Za nabavku materijala i to:")
    y -= 10 * mm

    # table layout
    col_w = [18*mm, 78*mm, 15*mm, 22*mm, 22*mm, 30*mm]
    headers = ["R.br.", "Ime proizvoda", "JM", "Količina", "Jed.cena", "Ukupno"]
    x0 = left_x
    row_h = 8 * mm

    def draw_row(cvs, ytop, values, bold=False):
        x = x0
        for i, val in enumerate(values):
            cvs.rect(x, ytop - row_h, col_w[i], row_h, stroke=1, fill=0)
            if bold:
                cvs.setFont(SERB_FONT_NAME, 11)
            else:
                cvs.setFont(SERB_FONT_NAME, 11)
            # alignment
            if i in (0, 2):
                cvs.drawCentredString(x + col_w[i]/2, ytop - 6*mm + 2, str(val))
            elif i in (3, 4, 5):
                cvs.drawRightString(x + col_w[i] - 2*mm, ytop - 6*mm + 2, str(val))
            else:
                cvs.drawString(x + 2*mm, ytop - 6*mm + 2, str(val))
            x += col_w[i]

    # header row
    draw_row(c, y, headers, bold=True)
    y -= row_h

    # lines
    for idx, r in enumerate(lines_df.to_dict("records"), start=1):
        uk = f"{_fmt_rs_money(_to_num(r.get('line_net',0)))} {currency}"
        draw_row(c, y, [
            idx,
            r.get("description",""),
            r.get("uom",""),
            _fmt_qty(_to_num(r.get("qty",0))),
            _fmt_rs_money(_to_num(r.get("unit_price",0))),
            uk
        ])
        y -= row_h
        if y < 40*mm:
            c.showPage()
            c.setFont(SERB_FONT_NAME, 12)
            y = h - 20*mm

    y -= 8 * mm

    # totals (right aligned)
    c.setFont(SERB_FONT_NAME, 12)
    c.drawRightString(w - 20*mm, y, f"UKUPNO : {_fmt_rs_money(totals['total_net'])} {currency}")
    y -= 7*mm
    c.drawRightString(w - 20*mm, y, f"OSNOVICA ZA OBRAČUN PDV: {_fmt_rs_money(totals['total_net'])} {currency}")
    y -= 7*mm
    c.drawRightString(w - 20*mm, y, f"PDV {int(vat_percent)}%: {_fmt_rs_money(totals['total_vat'])} {currency}")
    y -= 9*mm
    c.setFont(SERB_FONT_NAME, 12)
    c.drawRightString(w - 20*mm, y, f"SVE UKUPNO ZA UPLATU: {_fmt_rs_money(totals['total_gross'])} {currency}")
    y -= 12*mm

    # footer notes
    c.setFont(SERB_FONT_NAME, 11)
    c.drawString(left_x, y, f"Paritet isporuke: {company_block.get('delivery_parity') or '/'}")
    y -= 6*mm
    c.drawString(left_x, y, f"Napomena o poreskom oslobađanju: {company_block.get('tax_note') or '/'}")
    y -= 6*mm
    c.drawString(left_x, y, company_block.get("payment_note") or "")
    if meta.get("nbs_rate_services"):
        y -= 6*mm
        c.drawString(left_x, y, f"Za obra?un radova primenjen srednji kurs NBS na dan {meta.get("nbs_rate_services_date") or ""}. {meta.get("nbs_rate_services")}")
    if meta.get("nbs_rate_advance"):
        y -= 6*mm
        c.drawString(left_x, y, f"Za obra?un avansa primenjen srednji kurs NBS na dan {meta.get("nbs_rate_advance_date") or ""}. {meta.get("nbs_rate_advance")}")
    if meta.get("advance_applied_amount"):
        y -= 6*mm
        c.drawString(left_x, y, f"Umanjenje po avansu: {_fmt_rs_money(_to_num(meta.get("advance_applied_amount"),0.0))} {currency}")
    if meta.get("payment_purpose"):
        y -= 6*mm
        c.drawString(left_x, y, f"Svrha uplate: {meta.get('payment_purpose')}")
    if meta.get("invoice_note"):
        y -= 6*mm
        c.drawString(left_x, y, f"Napomena: {meta.get('invoice_note')}")
    y -= 10*mm
    c.drawRightString(w - 20*mm, y, f"■ {company_block.get('company_name') or 'LASER LUX'} ■ d.o.o.")

    c.showPage()
    c.save()
    buf.seek(0)
    return buf

def pdf_to_images(pdf_bytes: bytes, zoom: float = 2.0):
    if not HAVE_PYMUPDF:
        return []
    imgs = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    mat = fitz.Matrix(zoom, zoom)
    for i in range(doc.page_count):
        page = doc.load_page(i)
        pix = page.get_pixmap(matrix=mat)
        imgs.append(pix.tobytes("png"))
    return imgs

# =========================
# LOAD for obracun from DB
# =========================
def load_invoice_items_for_calc(invoice_id: int) -> pd.DataFrame:
    sql = """
    SELECT
      i.invoice_no AS "Broj računa",
      c.name AS "Kompanija",
      it.material_id AS "ID materijala",
      m.name AS "Materijal",
      it.qty AS "Količina za fakturisanje",
      it.uom AS "Jedinica mere za fakturisanje",
      m.uom AS "Jedinica mere za lager - skidanje količine",
      m.opis_materijala AS "Opis Materijala",
      m.tech_normative AS "Normativna potrošnja (tehnički list)",
      it.qty AS "Količina za fakturisanje (ono što piše u tabeli za račune - Normative)",
      it.uom AS "Jedinica mere za fakturisanje - u računu"
    FROM invoices i
    JOIN clients c ON c.client_id = i.client_id
    JOIN invoice_items it ON it.invoice_id = i.invoice_id
    JOIN materials m ON m.material_id = it.material_id
    WHERE i.invoice_id = :invoice_id
    ORDER BY it.item_id;
    """
    return db_df(sql, {"invoice_id": invoice_id})

# =========================
# UI
# =========================
tabs = st.tabs(["🧾 Novi račun", "📄 Računi", "🧮 Obračun + Word + Lager"])

# ======================================================
# TAB 1: Novi račun (Preview PDF + Save + Word)
# ======================================================
with tabs[0]:
    
    # =========================
    # TOP ROW: 3 columns (client edit / add client / company settings)
    # =========================
    st.markdown("### 🧩 Podešavanja (klijent / novi klijent / prodavac)")

    colL, colM, colR = st.columns(3)

    # ---------- 1) EDIT CLIENT ----------
    with colL:
        with st.expander("✏️ Uredi klijenta (iz baze)", expanded=False):

            clients = db_df("SELECT client_id, name FROM clients ORDER BY name")
            if clients.empty:
                st.warning("Nema klijenata u bazi.")
            else:
                edit_client_id = st.selectbox(
                    "Izaberi klijenta",
                    clients["client_id"].tolist(),
                    format_func=lambda x: clients.loc[clients["client_id"] == x, "name"].iloc[0],
                    key="edit_client_id_top"
                )

                crow = db_df("""
                    SELECT client_id, name, address, city, pib, mb, phone, email
                    FROM clients
                    WHERE client_id=:id
                """, {"id": int(edit_client_id)}).iloc[0].to_dict()

                with st.form("edit_client_form_top"):
                    name = st.text_input("Naziv", value=crow.get("name") or "")
                    c1, c2 = st.columns(2)
                    with c1:
                        address = st.text_input("Adresa", value=crow.get("address") or "")
                        city = st.text_input("Grad / Poštanski broj", value=crow.get("city") or "")
                        pib = st.text_input("PIB", value=crow.get("pib") or "")
                        mb = st.text_input("Matični broj", value=crow.get("mb") or "")
                    with c2:
                        phone = st.text_input("Telefon", value=crow.get("phone") or "")
                        email = st.text_input("E-mail", value=crow.get("email") or "")

                    save_edit = st.form_submit_button("💾 Sačuvaj izmene")

                if save_edit:
                    db_exec("""
                        UPDATE clients
                        SET name=:n, address=:a, city=:c, pib=:p, mb=:m, phone=:ph, email=:e
                        WHERE client_id=:id
                    """, {
                        "n": name.strip(),
                        "a": address.strip() or None,
                        "c": city.strip() or None,
                        "p": pib.strip() or None,
                        "m": mb.strip() or None,
                        "ph": phone.strip() or None,
                        "e": email.strip() or None,
                        "id": int(edit_client_id),
                    })
                    st.success("Klijent sačuvan ✅")
                    st.rerun()

    # ---------- 2) ADD NEW CLIENT ----------
    with colM:
        with st.expander("➕ Dodaj novog klijenta", expanded=False):
            with st.form("add_client_form_top"):
                n_name = st.text_input("Naziv*", value="")
                c1, c2 = st.columns(2)
                with c1:
                    n_address = st.text_input("Adresa", value="")
                    n_city = st.text_input("Grad / Poštanski broj", value="")
                    n_pib = st.text_input("PIB", value="")
                    n_mb = st.text_input("Matični broj", value="")
                with c2:
                    n_phone = st.text_input("Telefon", value="")
                    n_email = st.text_input("E-mail", value="")

                add_btn = st.form_submit_button("✅ Dodaj")

            if add_btn:
                if not n_name.strip():
                    st.error("Naziv je obavezan.")
                else:
                    db_exec("""
                        INSERT INTO clients(name, address, city, pib, mb, phone, email)
                        VALUES (:n, :a, :c, :p, :m, :ph, :e)
                    """, {
                        "n": n_name.strip(),
                        "a": n_address.strip() or None,
                        "c": n_city.strip() or None,
                        "p": n_pib.strip() or None,
                        "m": n_mb.strip() or None,
                        "ph": n_phone.strip() or None,
                        "e": n_email.strip() or None,
                    })
                    st.success("Klijent dodat ✅")
                    st.rerun()

    # ---------- 3) COMPANY SETTINGS / SELLER ----------
    with colR:
        with st.expander("🏬 Prodavac (podešavanja firme)", expanded=False):
            cs = db_df("SELECT * FROM company_settings ORDER BY id LIMIT 1")
            if cs.empty:
                db_exec("INSERT INTO company_settings(id, company_name) VALUES (1,'LASER - LUX d.o.o.')")
                cs = db_df("SELECT * FROM company_settings WHERE id=1")

            cs = cs.iloc[0].to_dict()

            with st.form("company_settings_form_top"):
                company_name = st.text_input("Naziv", value=cs.get("company_name") or "")
                address = st.text_input("Adresa", value=cs.get("address") or "")
                city = st.text_input("Grad / Poštanski broj", value=cs.get("city") or "")
                country = st.text_input("Država", value=cs.get("country") or "")
                phone = st.text_input("Telefon", value=cs.get("phone") or "")
                fax = st.text_input("Fax", value=cs.get("fax") or "")
                email = st.text_input("Email", value=cs.get("email") or "")
                bank_account = st.text_input("Tekući račun", value=cs.get("bank_account") or "")
                pib = st.text_input("PIB", value=cs.get("pib") or "")
                mb = st.text_input("Matični broj", value=cs.get("mb") or "")
                delivery_parity = st.text_input("Paritet isporuke", value=cs.get("delivery_parity") or "/")
                tax_note = st.text_input("Napomena o poreskom oslobađanju", value=cs.get("tax_note") or "/")
                payment_note = st.text_area(
                    "Napomena za uplatu",
                    value=cs.get("payment_note") or "Uplatu izvršiti na tekući račun br.: 205-508589-34 kod NLB komercijalna banka a.d.",
                    height=80
                )

                save_company = st.form_submit_button("💾 Sačuvaj firmu")

            if save_company:
                db_exec("""
                    UPDATE company_settings
                    SET company_name=:company_name,
                        address=:address,
                        city=:city,
                        country=:country,
                        phone=:phone,
                        fax=:fax,
                        email=:email,
                        bank_account=:bank_account,
                        pib=:pib,
                        mb=:mb,
                        delivery_parity=:delivery_parity,
                        tax_note=:tax_note,
                        payment_note=:payment_note
                    WHERE id=1
                """, {
                    "company_name": company_name.strip(),
                    "address": address.strip() or None,
                    "city": city.strip() or None,
                    "country": country.strip() or None,
                    "phone": phone.strip() or None,
                    "fax": fax.strip() or None,
                    "email": email.strip() or None,
                    "bank_account": bank_account.strip() or None,
                    "pib": pib.strip() or None,
                    "mb": mb.strip() or None,
                    "delivery_parity": delivery_parity.strip() or None,
                    "tax_note": tax_note.strip() or None,
                    "payment_note": payment_note.strip() or None,
                })
                st.success("Prodavac sačuvan ✅")
                st.rerun()

    st.markdown("---")

    st.subheader("🧾 Kreiraj račun (preview kao štampa + Word)")
    # ---- Company settings editor (vidljivo ovde)


    materials = db_df("SELECT material_id, name, uom FROM materials ORDER BY name")


    with st.expander("? Dodaj novi materijal", expanded=False):
        new_mat_name = st.text_input("Naziv materijala", value="")
        new_mat_uom = st.text_input("JM (lager)", value="")
        new_mat_desc = st.text_input("Opis materijala (opciono)", value="")
        new_mat_norm = st.text_input("Normativ (opciono)", value="")
        if st.button("Sa?uvaj materijal", key="add_material_btn"):
            if not new_mat_name.strip():
                st.error("Unesi naziv materijala.")
            else:
                try:
                    with engine.begin() as c:
                        try:
                            c.execute(text("""
                                INSERT INTO materials(name, uom, opis_materijala, tech_normative)
                                VALUES (:n, :u, :d, :t)
                            """), {
                                "n": new_mat_name.strip(),
                                "u": new_mat_uom.strip() or None,
                                "d": new_mat_desc.strip() or None,
                                "t": new_mat_norm.strip() or None,
                            })
                        except Exception:
                            c.execute(text("""
                                INSERT INTO materials(name, uom)
                                VALUES (:n, :u)
                            """), {
                                "n": new_mat_name.strip(),
                                "u": new_mat_uom.strip() or None,
                            })
                    st.success("Materijal dodat ?")
                    st.rerun()
                except Exception as e:
                    st.error(f"Gre?ka pri upisu: {e}")

    if clients.empty or materials.empty:
        st.warning("Nema klijenata ili materijala u bazi.")
    else:
        if "new_items" not in st.session_state:
            st.session_state["new_items"] = pd.DataFrame(
                columns=["material_id", "qty", "uom", "unit_price", "discount", "description"]
            )
        

        inv_type = st.selectbox("Tip računa", INVOICE_TYPES, key="inv_type_outside")

        with st.form("new_invoice_form", clear_on_submit=False):
            col1, col2, col3, col4 = st.columns([2, 1, 1, 1])

            with col1:
                client_id = st.selectbox(
                    "Klijent",
                    clients["client_id"].tolist(),
                    format_func=lambda x: clients.loc[clients["client_id"] == x, "name"].iloc[0],
                )

                client_row = db_df(
                    "SELECT client_id, name, address, city, pib, mb, phone, email FROM clients WHERE client_id=:id",
                    {"id": int(client_id)}
                ).iloc[0].to_dict()

            with col2:
                invoice_no = st.text_input("Broj ra?una")
            with col3:
                issue_date = st.date_input("Datum", value=date.today())
            with col4:
                delivery_note_no = st.text_input("Broj otpremnice", value="")

            r1, r2, r3, r4 = st.columns(4)
            with r1:
                vat_percent = st.number_input("PDV (%)", min_value=0.0, max_value=50.0, value=20.0, step=1.0)
            with r2:
                currency = st.selectbox("Valuta", ["RSD", "EUR"], index=0)
            with r3:
                due_date = st.date_input("Rok pla?anja", value=issue_date)
            with r4:
                place_of_issue = st.text_input("Mesto izdavanja", value="Beograd")

            r5, r6, r7, r8 = st.columns(4)
            with r5:
                place_of_supply = st.text_input("Mesto prometa", value="Beograd")
            with r6:
                supply_date = st.date_input("Datum prometa", value=issue_date)
            with r7:
                fx_rate = st.number_input("Kurs EUR (RSD/EUR)", min_value=0.0, value=0.0, step=0.0001)
            with r8:
                nbs_rate_services_date = st.text_input("Datum kursa NBS (radovi)", value="")

            r9, r10, r11, r12 = st.columns(4)
            with r9:
                nbs_rate_services = st.number_input("Kurs NBS za radove", min_value=0.0, value=0.0, step=0.0001)
            with r10:
                nbs_rate_advance_date = st.text_input("Datum kursa NBS (avans)", value="")
            with r11:
                nbs_rate_advance = st.number_input("Kurs NBS za avans", min_value=0.0, value=0.0, step=0.0001)
            with r12:
                adv_list = db_df("""
                    SELECT invoice_id, invoice_no, total_gross
                    FROM invoices
                    WHERE invoice_type IN ('ADVANCE','AVANSI')
                    ORDER BY invoice_id DESC
                    LIMIT 200
                """)
                adv_pick = st.selectbox(
                    "Pove?i avansni ra?un (opciono)",
                    adv_list["invoice_id"].tolist() if not adv_list.empty else [],
                    format_func=lambda x: f"{adv_list.loc[adv_list.invoice_id==x,'invoice_no'].iloc[0]} | {adv_list.loc[adv_list.invoice_id==x,'total_gross'].iloc[0]:.2f}" if x in adv_list["invoice_id"].values else str(x),
                    index=0 if not adv_list.empty else None,
                ) if not adv_list.empty else None

            r13, r14, r15, r16 = st.columns(4)
            with r13:
                advance_applied_amount = st.number_input("Iznos avansa za umanjenje", min_value=0.0, value=0.0, step=1.0)
            with r14:
                payment_purpose = st.text_input("Svrha uplate", value="", key="payment_purpose_input")
            with r15:
                invoice_note = st.text_input("Napomena", value="", key="invoice_note_input")
            with r16:
                st.markdown(" ")

            # defaults for privremena fields
            situation_title = ""
            situation_period = ""
            object_desc = ""
            contract_value_eur = 0.0
            works_place = ""
            base_contract_no = ""
            base_contract_date = ""
            add_contract_no = ""
            add_contract_date = ""
            cumulative_current_eur = 0.0
            cumulative_prev_eur = 0.0
            monthly_current_eur = 0.0
            monthly_current_rsd = 0.0
            advance_deduction_eur = 0.0
            advance_deduction_rsd = 0.0
            total_due_eur = 0.0
            total_due_rsd = 0.0
            vat_note_override = ""
            payment_currency_text = ""
            supervisor_name = ""
            receiver_name = ""
            receiver_title = ""
            supplier_name = ""
            supplier_title = ""

            inv_type_local = st.session_state.get("inv_type_outside", inv_type)

            # PRIVREMENA SITUACIJA polja (ručni unos)
            if _is_temp_situation(inv_type_local):
                st.markdown("### Privremena situacija")
                ps1, ps2, ps3, ps4 = st.columns(4)
                with ps1:
                    situation_title = st.text_input("Naziv situacije", value="PRIVREMENA SITUACIJA")
                    situation_period = st.text_input("Period", value="")
                    works_place = st.text_input("Mesto izvođenja radova", value="")
                with ps2:
                    contract_value_eur = st.number_input("Vrednost ugovorenih radova (EUR)", min_value=0.0, value=0.0, step=1.0)
                    base_contract_no = st.text_input("Broj osnovnog ugovora", value="")
                    base_contract_date = st.text_input("Datum osnovnog ugovora", value="")
                with ps3:
                    cumulative_current_eur = st.number_input("1. Kumulativ po tekućoj situaciji (EUR)", min_value=0.0, value=0.0, step=1.0)
                    cumulative_prev_eur = st.number_input("2. Kumulativ po prethodnoj situaciji (EUR)", min_value=0.0, value=0.0, step=1.0)
                    monthly_current_eur = st.number_input("3. Mesečno po tekućoj situaciji (EUR)", min_value=0.0, value=0.0, step=1.0)
                with ps4:
                    advance_deduction_eur = st.number_input("Umanjenje po avansu (EUR)", min_value=0.0, value=0.0, step=1.0)
                    total_due_eur = st.number_input("Ukupno za naplatu (EUR)", min_value=0.0, value=0.0, step=1.0)
                    payment_currency_text = st.text_input("Valuta plaćanja", value="Prema ugovoru")

                object_desc = st.text_area("Na objektu", value="", height=80)
                add_contract_no = st.text_input("Broj dopunskog ugovora", value="/")
                add_contract_date = st.text_input("Datum dopunskog ugovora", value="")
                vat_note_override = st.text_input("Napomena za PDV", value="PDV se ne obračunava na osnovu člana 10, stav 2, tačka 3, Zakona o PDVu.")

                # RSD preračun iz EUR * kurs
                monthly_current_rsd = round(float(monthly_current_eur or 0.0) * float(fx_rate or 0.0), 2)
                advance_deduction_rsd = round(float(advance_deduction_eur or 0.0) * float(fx_rate or 0.0), 2)
                total_due_rsd = round(float(total_due_eur or 0.0) * float(fx_rate or 0.0), 2)

                st.write(f"Mesečno (RSD): {monthly_current_rsd:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
                st.write(f"Umanjenje avansa (RSD): {advance_deduction_rsd:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
                st.write(f"Ukupno za naplatu (RSD): {total_due_rsd:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))

                s1, s2, s3, s4 = st.columns(4)
                with s1:
                    supervisor_name = st.text_input("Nadzorni organ", value="")
                with s2:
                    receiver_name = st.text_input("Primalac (ime)", value="")
                    receiver_title = st.text_input("Primalac (funkcija)", value="")
                with s3:
                    supplier_name = st.text_input("Ispručilac (ime)", value="")
                    supplier_title = st.text_input("Ispručilac (funkcija)", value="")
                with s4:
                    st.markdown(" ")
            advance_amount = None
            service_items = None

            if _is_advance(inv_type_local):
                st.markdown("### Avans")
                advance_amount = st.number_input("Iznos avansa", min_value=0.0, value=0.0, step=1.0)
            elif _is_service(inv_type_local):
                st.markdown("### Usluge")
                if "service_items" not in st.session_state:
                    st.session_state["service_items"] = pd.DataFrame(
                        columns=["material_id", "name", "item_note", "uom", "unit_price", "total_price"]
                    )
                service_items = st.data_editor(
                    st.session_state["service_items"],
                    num_rows="dynamic",
                    use_container_width=True,
                    key="service_items_editor",
                    column_config={
                        "material_id": st.column_config.SelectboxColumn(
                            "Ime proizvoda",
                            options=materials["material_id"].tolist(),
                            format_func=lambda x: materials.loc[materials["material_id"] == x, "name"].iloc[0]
                            if x in materials["material_id"].values else str(x),
                        ),
                        "name": st.column_config.TextColumn("Ime proizvoda (ručno)"),
                        "item_note": st.column_config.TextColumn("Opis"),
                        "uom": st.column_config.TextColumn("Jedinica mere"),
                        "unit_price": st.column_config.NumberColumn("Jedinicna cena"),
                        "total_price": st.column_config.NumberColumn("Ukupna cena"),
                    },
                )
                # autopopuni name/uom iz baze ako je izabran materijal
                for idx, row in service_items.iterrows():
                    mid = row.get("material_id")
                    if pd.notna(mid) and mid in materials["material_id"].values:
                        if pd.isna(row.get("name")) or str(row.get("name")).strip() == "":
                            service_items.at[idx, "name"] = materials.loc[materials["material_id"] == mid, "name"].iloc[0]
                        if pd.isna(row.get("uom")) or str(row.get("uom")).strip() == "":
                            service_items.at[idx, "uom"] = materials.loc[materials["material_id"] == mid, "uom"].iloc[0]
                st.session_state["service_items"] = service_items
            elif _is_temp_situation(inv_type_local):
                st.info("Privremena situacija nema stavke materijala.")
            else:
                st.markdown("### Stavke")

                df_new = st.data_editor(
                    st.session_state["new_items"],
                    num_rows="dynamic",
                    use_container_width=True,
                    key="new_items_editor",
                    column_config={
                        "material_id": st.column_config.SelectboxColumn(
                            "Materijal",
                            options=materials["material_id"].tolist(),
                            format_func=lambda x: materials.loc[materials["material_id"] == x, "name"].iloc[0]
                            if x in materials["material_id"].values else str(x),
                        ),
                        "qty": st.column_config.NumberColumn("Koli?ina"),
                        "uom": st.column_config.TextColumn("JM (fakt)"),
                        "unit_price": st.column_config.NumberColumn("Jed.cena"),
                        "discount": st.column_config.NumberColumn("Popust (RSD)"),
                        "description": st.column_config.TextColumn("Opis (opciono)"),
                    },
                )

                # autopopuni
                for idx, row in df_new.iterrows():
                    mid = row.get("material_id")
                    if pd.notna(mid) and mid in materials["material_id"].values:
                        if pd.isna(row.get("uom")) or str(row.get("uom")).strip() == "":
                            df_new.at[idx, "uom"] = materials.loc[materials["material_id"] == mid, "uom"].iloc[0]
                        if pd.isna(row.get("description")) or str(row.get("description")).strip() == "":
                            df_new.at[idx, "description"] = materials.loc[materials["material_id"] == mid, "name"].iloc[0]
                    if pd.isna(row.get("discount")):
                        df_new.at[idx, "discount"] = 0

                st.session_state["new_items"] = df_new

            do_preview = st.form_submit_button("👀 Preview (kao štampa)")
            do_save = st.form_submit_button("💾 Sačuvaj u bazu")



        

        # Preview PDF (u expanderu) + download PDF
        if do_preview:
            inv_type_local = st.session_state.get("inv_type_outside", inv_type)
            if _is_temp_situation(inv_type_local):
                work = pd.DataFrame()
                lines = pd.DataFrame()
                totals = {
                    "total_net": float(total_due_eur or 0.0),
                    "total_vat": 0.0,
                    "total_gross": float(total_due_eur or 0.0),
                }
            elif _is_advance(inv_type_local):
                if advance_amount is None or float(advance_amount) <= 0:
                    st.error("Unesi iznos avansa.")
                    work = pd.DataFrame()
                else:
                    mid = _get_or_create_advance_material_id()
                    work = pd.DataFrame([{
                        "material_id": mid,
                        "qty": 1,
                        "uom": currency,
                        "unit_price": float(advance_amount),
                        "discount": 0.0,
                        "description": "AVANS"
                    }])
            elif _is_service(inv_type_local):
                work = st.session_state.get("service_items", pd.DataFrame()).copy()
                work = work[work["name"].notna()].copy() if not work.empty else work
                if work.empty:
                    st.error("Dodaj bar jednu stavku usluge.")
                else:
                    work = build_service_lines(work)
            else:
                work = df_new.copy()
                work = work[work["material_id"].notna() & work["qty"].notna()].copy()
                if work.empty:
                    st.error("Dodaj bar jednu stavku (materijal + koli?ina).")

            if _is_temp_situation(inv_type_local):
                pass
            elif work.empty:
                pass
            else:
                vat_rate = float(vat_percent) / 100.0
                work["unit_price"] = work.get("unit_price", 0).fillna(0)
                work["discount"] = work.get("discount", 0).fillna(0)
                lines = calc_lines(work, vat_rate)
                totals = sum_totals(lines)

                cs = db_df("SELECT * FROM company_settings WHERE id=1").iloc[0].to_dict()
                cli = db_df("SELECT name, address, city FROM clients WHERE client_id=:id", {"id": int(client_id)}).iloc[0].to_dict()

                meta = {
                    "invoice_no": str(invoice_no).strip(),
                    "issue_date_str": issue_date.strftime("%d.%m.%Y."),
                    "due_date_str": due_date.strftime("%d.%m.%Y.") if due_date else "/",
                    "place_of_issue": place_of_issue,
                    "delivery_note_no": delivery_note_no.strip() or None,
                    "place_of_supply": place_of_supply,
                    "supply_date_str": supply_date.strftime("%d.%m.%Y.") if supply_date else "",

                    
                    "invoice_type": inv_type_local,
                    "fx_rate": fx_rate,
                    "nbs_rate_services": nbs_rate_services,
                    "nbs_rate_services_date": nbs_rate_services_date,
                    "nbs_rate_advance": nbs_rate_advance,
                    "nbs_rate_advance_date": nbs_rate_advance_date,
                    "advance_applied_amount": advance_applied_amount,
                    "situation_title": situation_title,
                    "situation_period": situation_period,
                    "object_desc": object_desc,
                    "contract_value_eur": contract_value_eur,
                    "works_place": works_place,
                    "base_contract_no": base_contract_no,
                    "base_contract_date": base_contract_date,
                    "add_contract_no": add_contract_no,
                    "add_contract_date": add_contract_date,
                    "cumulative_current_eur": cumulative_current_eur,
                    "cumulative_prev_eur": cumulative_prev_eur,
                    "monthly_current_eur": monthly_current_eur,
                    "monthly_current_rsd": monthly_current_rsd,
                    "advance_deduction_eur": advance_deduction_eur,
                    "advance_deduction_rsd": advance_deduction_rsd,
                    "total_due_eur": total_due_eur,
                    "total_due_rsd": total_due_rsd,
                    "vat_note_override": vat_note_override,
                    "payment_currency_text": payment_currency_text,
                    "supervisor_name": supervisor_name,
                    "receiver_name": receiver_name,
                    "receiver_title": receiver_title,
                    "supplier_name": supplier_name,
                    "supplier_title": supplier_title,
                    "payment_purpose": payment_purpose.strip() or None,
                    "invoice_note": invoice_note.strip() or None,
                }

                pdf_buf = build_invoice_pdf_preview(
                    company_block=cs,
                    client_block=cli,
                    meta=meta,
                    lines_df=lines,
                    totals=totals,
                    currency=currency,
                    vat_percent=float(vat_percent)
                )
                st.session_state["last_preview_pdf"] = pdf_buf.getvalue()

                with st.expander("🖨️ Preview (kao štampa) — PDF", expanded=True):
                    if HAVE_PYMUPDF:
                        imgs = pdf_to_images(st.session_state["last_preview_pdf"], zoom=2.0)
                        if imgs:
                            for i, img in enumerate(imgs, start=1):
                                st.image(img, caption=f"Strana {i}", use_container_width=True)
                        else:
                            st.info("Nisam uspeo da renderujem slike iz PDF-a. Preuzmi PDF ispod.")
                    else:
                        st.warning("Za prikaz kao slike instaliraj: pip install pymupdf")

                    st.download_button(
                        "⬇️ Preuzmi PDF preview",
                        data=st.session_state["last_preview_pdf"],
                        file_name=f"preview_racun_{meta['invoice_no'] or 'draft'}.pdf",
                        mime="application/pdf",
                    )

        # Save in DB + allow Word generation
        if do_save:
            if not str(invoice_no).strip():
                st.error("Unesi broj računa.")
            else:
                inv_type_local = st.session_state.get("inv_type_outside", inv_type)
                if _is_temp_situation(inv_type_local):
                    work = pd.DataFrame()
                    lines = pd.DataFrame()
                    totals = {
                        "total_net": float(total_due_eur or 0.0),
                        "total_vat": 0.0,
                        "total_gross": float(total_due_eur or 0.0),
                    }
                elif _is_advance(inv_type_local):
                    if advance_amount is None or float(advance_amount) <= 0:
                        st.error("Unesi iznos avansa.")
                        work = pd.DataFrame()
                    else:
                        mid = _get_or_create_advance_material_id()
                        work = pd.DataFrame([{
                            "material_id": mid,
                            "qty": 1,
                            "uom": currency,
                            "unit_price": float(advance_amount),
                            "discount": 0.0,
                            "description": "AVANS"
                        }])
                elif _is_service(inv_type_local):
                    work = st.session_state.get("service_items", pd.DataFrame()).copy()
                    work = work[work["name"].notna()].copy() if not work.empty else work
                    if work.empty:
                        st.error("Dodaj bar jednu stavku usluge.")
                    else:
                        work = build_service_lines(work)
                else:
                    work = df_new.copy()
                    work = work[work["material_id"].notna() & work["qty"].notna()].copy()
                    if work.empty:
                        st.error("Dodaj bar jednu stavku (materijal + koli?ina).")

                if _is_temp_situation(inv_type_local):
                    pass
                elif work.empty:
                    pass
                else:
                    vat_rate = float(vat_percent) / 100.0
                    work["unit_price"] = work.get("unit_price", 0).fillna(0)
                    work["discount"] = work.get("discount", 0).fillna(0)
                    lines = calc_lines(work, vat_rate)
                    totals = sum_totals(lines)

                    with engine.begin() as c:
                        inv_id = c.execute(text("""
                            INSERT INTO invoices(
                                invoice_no, invoice_type, status, client_id,
                                issue_date, due_date, currency, vat_rate,
                                total_net, total_vat, total_gross,
                                delivery_note_no, place_of_issue, place_of_supply, supply_date,
                                payment_purpose, invoice_note,
                                fx_rate, nbs_rate_services, nbs_rate_services_date,
                                nbs_rate_advance, nbs_rate_advance_date,
                                advance_invoice_id, advance_applied_amount, total_due,
                                situation_title, situation_period, object_desc, contract_value_eur,
                                works_place, base_contract_no, base_contract_date,
                                add_contract_no, add_contract_date,
                                cumulative_current_eur, cumulative_prev_eur,
                                monthly_current_eur, monthly_current_rsd,
                                advance_deduction_eur, advance_deduction_rsd,
                                total_due_eur, total_due_rsd,
                                vat_note_override, payment_currency_text,
                                supervisor_name, receiver_name, receiver_title,
                                supplier_name, supplier_title
                            )
                            VALUES (
                                :no, :t, 'DRAFT', :cid,
                                :issue, :due, :cur, :vr,
                                :tn, :tv, :tg,
                                :dn, :pio, :ps, :sd,
                                :pp, :note,
                                :fx, :nbs_s, :nbs_sd,
                                :nbs_a, :nbs_ad,
                                :adv_id, :adv_amt, :due,
                                :sit_title, :sit_period, :obj_desc, :contract_eur,
                                :works_place, :base_no, :base_date,
                                :add_no, :add_date,
                                :cum_cur, :cum_prev,
                                :mon_cur, :mon_cur_rsd,
                                :adv_ded, :adv_ded_rsd,
                                :due_eur, :due_rsd,
                                :vat_note, :pay_curr,
                                :supervisor, :recv_name, :recv_title,
                                :supp_name, :supp_title
                            )
                            RETURNING invoice_id
                        """), {
                            "no": str(invoice_no).strip(),
                            "t": inv_type,
                            "cid": int(client_id),
                            "issue": issue_date,
                            "due": due_date,
                            "cur": currency,
                            "vr": vat_rate,
                            "tn": totals["total_net"],
                            "tv": totals["total_vat"],
                            "tg": totals["total_gross"],
                            "dn": delivery_note_no.strip() or None,
                            "pio": place_of_issue.strip() or None,
                            "ps": place_of_supply.strip() or None,
                            "sd": supply_date,
                            "pp": payment_purpose.strip() or None,
                            "note": invoice_note.strip() or None,
                            "fx": fx_rate or None,
                            "nbs_s": nbs_rate_services or None,
                            "nbs_sd": nbs_rate_services_date.strip() or None,
                            "nbs_a": nbs_rate_advance or None,
                            "nbs_ad": nbs_rate_advance_date.strip() or None,
                            "adv_id": int(adv_pick) if adv_pick else None,
                            "adv_amt": float(advance_applied_amount or 0.0),
                            "due": None,
                            "sit_title": situation_title or None,
                            "sit_period": situation_period or None,
                            "obj_desc": object_desc or None,
                            "contract_eur": float(contract_value_eur or 0.0),
                            "works_place": works_place or None,
                            "base_no": base_contract_no or None,
                            "base_date": base_contract_date or None,
                            "add_no": add_contract_no or None,
                            "add_date": add_contract_date or None,
                            "cum_cur": float(cumulative_current_eur or 0.0),
                            "cum_prev": float(cumulative_prev_eur or 0.0),
                            "mon_cur": float(monthly_current_eur or 0.0),
                            "mon_cur_rsd": float(monthly_current_rsd or 0.0),
                            "adv_ded": float(advance_deduction_eur or 0.0),
                            "adv_ded_rsd": float(advance_deduction_rsd or 0.0),
                            "due_eur": float(total_due_eur or 0.0),
                            "due_rsd": float(total_due_rsd or 0.0),
                            "vat_note": vat_note_override or None,
                            "pay_curr": payment_currency_text or None,
                            "supervisor": supervisor_name or None,
                            "recv_name": receiver_name or None,
                            "recv_title": receiver_title or None,
                            "supp_name": supplier_name or None,
                            "supp_title": supplier_title or None,
                        }).scalar()

                        for _, r in lines.iterrows():
                            c.execute(text("""
                                INSERT INTO invoice_items(
                                    invoice_id, material_id, description, item_note, qty, uom,
                                    unit_price, discount,
                                    line_net, line_vat, line_gross
                                )
                                VALUES (
                                    :iid, :mid, :desc, :note, :qty, :uom,
                                    :up, :disc,
                                    :ln, :lv, :lg
                                )
                            """), {
                                "iid": int(inv_id),
                                "mid": int(r["material_id"]),
                                "desc": str(r.get("description") or r.get("name") or ""),
                                "note": str(r.get("item_note") or ""),
                                "qty": float(_to_num(r.get("qty"), 0.0)),
                                "uom": str(r.get("uom") or ""),
                                "up": float(_to_num(r.get("unit_price"), 0.0)),
                                "disc": float(_to_num(r.get("discount"), 0.0)),
                                "ln": float(r["line_net"]),
                                "lv": float(r["line_vat"]),
                                "lg": float(r["line_gross"]),
                            })

                    st.success(f"Sačuvano ✅ (invoice_id={int(inv_id)})")
                    st.session_state["last_created_invoice_id"] = int(inv_id)

        # Word (račun) dugme u ovom tabu
        if st.session_state.get("last_created_invoice_id"):
            st.markdown("---")
            st.subheader("📄 Generiši Word račun (1:1)")

            inv_id = int(st.session_state["last_created_invoice_id"])

            c1, c2 = st.columns([1, 2])
            with c1:
                if st.button("📄 Generiši Word račun", key="gen_word_invoice_btn"):
                    try:
                        buf = generate_invoice_word_from_db(inv_id)
                        st.session_state["last_invoice_word"] = buf.getvalue()
                        st.success("Word je generisan ✅")
                    except Exception as e:
                        st.error(f"Greška pri generisanju Word-a: {e}")

            with c2:
                if st.session_state.get("last_invoice_word"):
                    st.download_button(
                        "⬇️ Preuzmi Word račun",
                        data=st.session_state["last_invoice_word"],
                        file_name=f"racun_{inv_id}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

# ======================================================
# TAB 2: Računi
# ======================================================
with tabs[1]:
    st.subheader("📄 Lista računa")

    inv = db_df("""
        SELECT i.invoice_id, i.invoice_no, i.invoice_type, i.status, i.issue_date, i.currency,
               COALESCE(i.total_net,0) AS total_net,
               COALESCE(i.total_vat,0) AS total_vat,
               COALESCE(i.total_gross,0) AS total_gross,
               c.name AS client
        FROM invoices i
        JOIN clients c ON c.client_id=i.client_id
        ORDER BY i.issue_date DESC NULLS LAST, i.invoice_id DESC
        LIMIT 1000
    """)
    st.dataframe(inv, use_container_width=True)

    pick = st.selectbox(
        "Detalji računa",
        inv["invoice_id"].tolist() if not inv.empty else [],
        format_func=lambda x: f"{inv.loc[inv.invoice_id==x,'invoice_no'].iloc[0]} — {inv.loc[inv.invoice_id==x,'client'].iloc[0]}",
    )

    if pick:
        items = db_df("""
            SELECT item_id, material_id, description, qty, uom,
                   COALESCE(unit_price,0) AS unit_price,
                   COALESCE(discount,0) AS discount,
                   COALESCE(line_net,0) AS line_net,
                   COALESCE(line_vat,0) AS line_vat,
                   COALESCE(line_gross,0) AS line_gross
            FROM invoice_items
            WHERE invoice_id = :id
            ORDER BY item_id
        """, {"id": int(pick)})

        st.markdown("### Stavke")
        st.dataframe(items, use_container_width=True)

        meta = inv.loc[inv.invoice_id == pick].iloc[0]
        st.info(
            f"NET: {meta['total_net']:.2f} {meta['currency']} | "
            f"PDV: {meta['total_vat']:.2f} {meta['currency']} | "
            f"BRUTO: {meta['total_gross']:.2f} {meta['currency']}"
        )

# ======================================================
# TAB 3: Obračun + Word + Lager (bulk u expanderu + search + rules u tab-u)
# ======================================================
with tabs[2]:
    st.subheader("🧮 Obračun + Word + Lager (DB)")

    # 2 kolone: levo obracun, desno pravila
    left_col, right_col = st.columns([2, 1])

    with right_col:
        st.markdown("### 🧮 Pravila (koeficijenti) — samo za Obračun tab")

        if "rules_df_calc" not in st.session_state:
            st.session_state["rules_df_calc"] = rules_to_df(material_rules)

        edited_rules_df = st.data_editor(
            st.session_state["rules_df_calc"],
            use_container_width=True,
            num_rows="dynamic",
            height=420,
            column_config={
                "rule_type": st.column_config.SelectboxColumn("rule_type", options=RULE_TYPES),
                "enabled": st.column_config.CheckboxColumn("enabled"),
                "factor": st.column_config.NumberColumn("factor"),
                "extra": st.column_config.NumberColumn("extra"),
            },
            key="rules_editor_calc_only",
        )

        if st.button("💾 Sačuvaj pravila (Obračun)", key="save_rules_calc_btn"):
            st.session_state["rules_df_calc"] = edited_rules_df
            st.success("Sačuvano ✅")

    with left_col:
        inv_pick = db_df("""
            SELECT i.invoice_id, i.invoice_no, i.invoice_type, i.status, i.stock_posted, c.name AS client
            FROM invoices i
            JOIN clients c ON c.client_id=i.client_id
            ORDER BY i.invoice_id DESC
            LIMIT 2000
        """)

        if inv_pick.empty:
            st.info("Nema računa u bazi.")
        else:
            with st.expander("📦 Bulk obračun (checkbox + select all) — preporučeno za DRAFT", expanded=False):
                show_draft_only = st.checkbox("Prikaži samo DRAFT", value=True, key="draft_only_bulk")
                q = st.text_input("🔎 Pretraga (broj računa ili klijent)", value="", key="bulk_search")

                view_df = inv_pick.copy()
                if show_draft_only:
                    view_df = view_df[view_df["status"] == "DRAFT"].copy()
                if q.strip():
                    qq = q.strip().lower()
                    view_df = view_df[
                        view_df["invoice_no"].astype(str).str.lower().str.contains(qq, na=False)
                        | view_df["client"].astype(str).str.lower().str.contains(qq, na=False)
                    ].copy()

                view_df = view_df.head(300)  # UI limit da ne ubije stranicu

                if "bulk_sel" not in st.session_state:
                    st.session_state["bulk_sel"] = {}

                for iid in view_df["invoice_id"].tolist():
                    st.session_state["bulk_sel"].setdefault(int(iid), False)

                b1, b2 = st.columns([1, 1])
                with b1:
                    if st.button("✅ Select all", key="bulk_all_btn"):
                        for iid in view_df["invoice_id"].tolist():
                            st.session_state["bulk_sel"][int(iid)] = True
                        st.rerun()
                with b2:
                    if st.button("⬜ Select none", key="bulk_none_btn"):
                        for iid in view_df["invoice_id"].tolist():
                            st.session_state["bulk_sel"][int(iid)] = False
                        st.rerun()

                sel_rows = []
                for _, r in view_df.iterrows():
                    iid = int(r["invoice_id"])
                    label = f'{r["invoice_no"]} — {r["client"]} ({r["invoice_type"]}, {r["status"]})'
                    checked = st.checkbox(label, value=st.session_state["bulk_sel"][iid], key=f"chk_{iid}")
                    st.session_state["bulk_sel"][iid] = bool(checked)
                    if checked:
                        sel_rows.append(iid)

                st.write(f"Izabrano računa: **{len(sel_rows)}**")

                colX, colY = st.columns([1, 1])
                with colX:
                    if st.button("🧮 Pokreni obračun za izabrane", key="bulk_run_btn"):
                        if not sel_rows:
                            st.error("Nisi izabrao nijedan račun.")
                        else:
                            results = []
                            for iid in sel_rows:
                                try:
                                    df_items = load_invoice_items_for_calc(int(iid))
                                    df_posle, _ = procesiraj_obracun_iz_db(
                                        df_items,
                                        edited_rules_df=st.session_state.get("rules_df_calc")
                                    )
                                    st.session_state.setdefault("bulk_df_posle", {})
                                    st.session_state["bulk_df_posle"][int(iid)] = df_posle
                                    results.append((iid, "OK", len(df_posle)))
                                except Exception as e:
                                    results.append((iid, f"ERR: {e}", 0))

                            st.success("Bulk obračun završen ✅")
                            st.dataframe(pd.DataFrame(results, columns=["invoice_id", "status", "rows"]), use_container_width=True)

                with colY:
                    if st.button("💾 Upiši invoice_allocations za izabrane", key="bulk_write_btn"):
                        if not sel_rows:
                            st.error("Nisi izabrao nijedan račun.")
                        else:
                            if "bulk_df_posle" not in st.session_state:
                                st.error("Prvo pokreni obračun za izabrane.")
                            else:
                                ok, err = 0, 0
                                for iid in sel_rows:
                                    df_posle = st.session_state["bulk_df_posle"].get(int(iid))
                                    if df_posle is None:
                                        err += 1
                                        continue

                                    try:
                                        db_exec("DELETE FROM invoice_allocations WHERE invoice_id=:id", {"id": int(iid)})
                                        with engine.begin() as c:
                                            for _, rr in df_posle.iterrows():
                                                mid = rr.get("ID materijala")
                                                if pd.isna(mid):
                                                    continue
                                                qtyw = rr.get("Kolicina za skidanje sa uracunatim Koef_novi za ovaj materijal")
                                                uomw = rr.get("Jedinica mere za lager - skidanje količine")
                                                note = rr.get("Napomena konverzije", "")

                                                c.execute(text("""
                                                    INSERT INTO invoice_allocations(invoice_id, material_id, qty_from_warehouse, uom, note)
                                                    VALUES (:iid, :mid, :q, :u, :n)
                                                """), {
                                                    "iid": int(iid),
                                                    "mid": int(float(mid)),
                                                    "q": float(qtyw) if pd.notna(qtyw) else 0.0,
                                                    "u": str(uomw) if pd.notna(uomw) else None,
                                                    "n": str(note)[:1000],
                                                })
                                        ok += 1
                                    except Exception:
                                        err += 1

                                st.success(f"Upisano ✅ OK={ok}, ERR={err}")

            st.markdown("---")
            st.markdown("### Pojedinačni rad (1 račun)")

            inv_id = st.selectbox(
                "Izaberi račun",
                inv_pick["invoice_id"].tolist(),
                format_func=lambda x: (
                    f'{inv_pick.loc[inv_pick["invoice_id"]==x, "invoice_no"].iloc[0]} — '
                    f'{inv_pick.loc[inv_pick["invoice_id"]==x, "client"].iloc[0]} '
                    f'({inv_pick.loc[inv_pick["invoice_id"]==x, "invoice_type"].iloc[0]}, '
                    f'{inv_pick.loc[inv_pick["invoice_id"]==x, "status"].iloc[0]})'
                ),
            )

            meta = inv_pick.loc[inv_pick["invoice_id"] == inv_id].iloc[0]
            is_final = str(meta["invoice_type"]).upper() in ("FINAL","PRODAJA MATERIJALA")
            posted = bool(meta.get("stock_posted"))

            cA, cB, cC = st.columns([1, 1, 2])

            with cA:
                if st.button("🧮 Pokreni obračun (1 račun)", key="one_calc_btn"):
                    df_items = load_invoice_items_for_calc(int(inv_id))
                    df_posle, _ = procesiraj_obracun_iz_db(
                        df_items,
                        edited_rules_df=st.session_state.get("rules_df_calc")
                    )
                    st.session_state["df_posle_db"] = df_posle
                    st.success("Obračun završen ✅")

            with cB:
                if st.button("💾 Upisi obračun u bazu (invoice_allocations)", key="one_write_alloc_btn"):
                    if "df_posle_db" not in st.session_state:
                        st.error("Prvo pokreni obračun.")
                    else:
                        df_posle = st.session_state["df_posle_db"]
                        db_exec("DELETE FROM invoice_allocations WHERE invoice_id=:id", {"id": int(inv_id)})

                        with engine.begin() as c:
                            for _, rr in df_posle.iterrows():
                                mid = rr.get("ID materijala")
                                if pd.isna(mid):
                                    continue
                                qtyw = rr.get("Kolicina za skidanje sa uracunatim Koef_novi za ovaj materijal")
                                uomw = rr.get("Jedinica mere za lager - skidanje količine")
                                note = rr.get("Napomena konverzije", "")

                                c.execute(text("""
                                    INSERT INTO invoice_allocations(invoice_id, material_id, qty_from_warehouse, uom, note)
                                    VALUES (:iid, :mid, :q, :u, :n)
                                """), {
                                    "iid": int(inv_id),
                                    "mid": int(float(mid)),
                                    "q": float(qtyw) if pd.notna(qtyw) else 0.0,
                                    "u": str(uomw) if pd.notna(uomw) else None,
                                    "n": str(note)[:1000],
                                })

                        st.success("Upisano ✅")

            with cC:
                if "df_posle_db" in st.session_state:
                    df_posle = st.session_state["df_posle_db"]
                    st.dataframe(df_posle, use_container_width=True)

                    broj = str(df_posle["Broj računa"].iloc[0])
                    word_buf = generate_word_for_racun(df_posle, broj)
                    st.download_button(
                        "📄 Preuzmi Word (obračun)",
                        data=word_buf,
                        file_name=f"obracun_{broj}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

            st.markdown("---")

            cnt = db_df("SELECT count(*) AS n FROM invoice_allocations WHERE invoice_id = :id", {"id": int(inv_id)})
            has_alloc = int(cnt.iloc[0]["n"]) > 0

            if is_final:
                if posted:
                    st.info("📦 Lager je već proknjižen za ovaj FINAL račun.")
                else:
                    if st.button("📦 Skini sa lagera (FINAL POST)", type="primary", key="final_post_btn"):
                        if not has_alloc:
                            st.error("Prvo upiši obračun u invoice_allocations.")
                        else:
                            try:
                                db_exec("SELECT post_invoice_stock(:id)", {"id": int(inv_id)})
                                st.success("Lager uspešno ažuriran ✅")
                            except Exception as e:
                                st.error(f"Greška pri knjiženju lagera: {e}")
            else:
                st.warning("Skidanje sa lagera je dozvoljeno samo za FINAL račun.")
